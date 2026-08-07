import streamlit as st
import sqlite3
import pandas as pd
from datetime import date
import hashlib
import secrets
import smtplib
import json
import os
import base64
import requests
from calendar import monthrange
from datetime import datetime
from email.mime.text import MIMEText
from io import BytesIO
from reportlab.platypus import (
    SimpleDocTemplate,
    Table,
    TableStyle,
    Paragraph,
    Spacer
)
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
import io
import tempfile
from docx import Document
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
import smtplib
from email.message import EmailMessage
from PyPDF2 import PdfReader, PdfWriter


# ============================================================
# CONFIGURAZIONE
# ============================================================

st.set_page_config(
    page_title="Statino Corsi Nuoto",
    page_icon="🏊",
    layout="wide"
)

DB_NAME = "corsi_nuoto.db"

BACKUP_ABILITATO = False

# ============================================================
# DATABASE
# ============================================================

conn = sqlite3.connect(
    DB_NAME,
    check_same_thread=False
)

c = conn.cursor()

# ============================================================
# FUNZIONI PASSWORD
# ============================================================
def crea_backup():

    backup = {}

    tabelle = [
        "utenti",
        "corsi",
        "corso_giorni",
        "bambini",
        "genitori_bambini",
        "assegnazioni_istruttori",
        "presenze",
        "stagioni",
        "sistema",
        "chiusure"
    ]

    for tabella in tabelle:

        try:

            df = pd.read_sql(
                f"SELECT * FROM {tabella}",
                conn
            )

            backup[tabella] = df.to_dict(
                orient="records"
            )

        except:

            pass

    with open(
        "backup_completo.json",
        "w",
        encoding="utf-8"
    ) as f:

        json.dump(
            backup,
            f,
            ensure_ascii=False,
            indent=4
        )

def elimina_istruttore(istruttore_id):

    c.execute(
        """
        DELETE FROM assegnazioni_istruttori
        WHERE istruttore_id = ?
        """,
        (istruttore_id,)
    )

    c.execute(
        """
        DELETE FROM utenti
        WHERE id = ?
        AND ruolo = 'istruttore'
        """,
        (istruttore_id,)
    )

    conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )
    
    except Exception as e:
    
        print(
            f"Errore backup GitHub: {e}"
        )

def calcola_eta(data_nascita):

    oggi = date.today()

    eta = (
        oggi.year
        - data_nascita.year
        - (
            (oggi.month, oggi.day)
            <
            (data_nascita.month, data_nascita.day)
        )
    )

    return eta
    
def elimina_manager(manager_id):

    c.execute(
        """
        DELETE FROM utenti
        WHERE id = ?
        AND ruolo = 'manager'
        """,
        (manager_id,)
    )

    conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )

    except:
        pass
        
def ripristina_backup():

    with open(
        "backup_completo.json",
        "r",
        encoding="utf-8"
    ) as f:

        backup = json.load(f)

    ordine = [
        "presenze",
        "assegnazioni_istruttori",
        "bambini",
        "corso_giorni",
        "corsi",
        "utenti",
        "stagioni",
        "chiusure"
    ]

    for tabella in ordine:

        try:

            c.execute(
                f"DELETE FROM {tabella}"
            )

        except:

            pass

    conn.commit()

    for tabella, records in backup.items():

        if len(records) == 0:
            continue

        df = pd.DataFrame(records)

        df.to_sql(
            tabella,
            conn,
            if_exists="append",
            index=False
        )

def hash_password(password, salt=None):

    if salt is None:
        salt = secrets.token_hex(16)

    password_hash = hashlib.sha256(
        (password + salt).encode("utf-8")
    ).hexdigest()

    return password_hash, salt


def verifica_password(password_inserita, password_hash, salt):

    nuovo_hash, _ = hash_password(
        password_inserita,
        salt
    )

    return nuovo_hash == password_hash

def genera_password_casuale():

    return secrets.token_urlsafe(9)
    
# ============================================================
# TABELLE
# ============================================================

c.execute("""
CREATE TABLE IF NOT EXISTS utenti (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    username TEXT UNIQUE NOT NULL,
    nome TEXT NOT NULL,
    ruolo TEXT NOT NULL,
    password_hash TEXT NOT NULL,
    salt TEXT NOT NULL,
    attivo INTEGER DEFAULT 1
)
""")

try:

    c.execute("""
    ALTER TABLE utenti
    ADD COLUMN password_visibile TEXT
    """)

    conn.commit()

except Exception:
    pass

c.execute("""
CREATE TABLE IF NOT EXISTS corsi (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    nome TEXT NOT NULL,
    livello TEXT,
    giorno TEXT,
    orario TEXT,
    stagione TEXT,
    attivo INTEGER DEFAULT 1
)
""")

c.execute("""
CREATE TABLE IF NOT EXISTS corso_giorni (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    corso_id INTEGER NOT NULL,
    giorno TEXT NOT NULL,
    orario TEXT NOT NULL,
    ordine INTEGER DEFAULT 1,
    FOREIGN KEY(corso_id) REFERENCES corsi(id)
)
""")

c.execute("""
CREATE TABLE IF NOT EXISTS bambini (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    nome TEXT NOT NULL,
    cognome TEXT NOT NULL,
    corso_id INTEGER,
    email_genitore TEXT,
    telefono_genitore TEXT,
    note TEXT,
    attivo INTEGER DEFAULT 1
)
""")
try:

    c.execute("""
        ALTER TABLE bambini
        ADD COLUMN corso_id INTEGER
    """)

    conn.commit()

except Exception:
    pass

try:

    c.execute("""
        ALTER TABLE bambini
        ADD COLUMN email_genitore TEXT
    """)

    conn.commit()

except:
    pass

try:

    c.execute("""
        ALTER TABLE bambini
        ADD COLUMN telefono_genitore TEXT
    """)

    conn.commit()

except:
    pass
    
c.execute("""
CREATE TABLE IF NOT EXISTS assegnazioni_istruttori (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    istruttore_id INTEGER NOT NULL,
    corso_id INTEGER NOT NULL,
    data_specifica TEXT,
    attiva INTEGER DEFAULT 1,
    FOREIGN KEY(istruttore_id) REFERENCES utenti(id),
    FOREIGN KEY(corso_id) REFERENCES corsi(id)
)
""")

c.execute("""
CREATE TABLE IF NOT EXISTS presenze (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    bambino_id INTEGER NOT NULL,
    corso_id INTEGER NOT NULL,
    data TEXT NOT NULL,
    presenza INTEGER NOT NULL,
    note TEXT,
    inserito_da INTEGER,
    UNIQUE(
        bambino_id,
        corso_id,
        data
    ),
    FOREIGN KEY(bambino_id) REFERENCES bambini(id),
    FOREIGN KEY(corso_id) REFERENCES corsi(id),
    FOREIGN KEY(inserito_da) REFERENCES utenti(id)
)
""")

c.execute("""
CREATE TABLE IF NOT EXISTS stagioni (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    nome TEXT UNIQUE NOT NULL,
    attiva INTEGER DEFAULT 1
)
""")

try:

    c.execute(
        """
        ALTER TABLE stagioni
        ADD COLUMN attiva INTEGER DEFAULT 1
        """
    )

    conn.commit()

except:

    pass

c.execute("""
CREATE TABLE IF NOT EXISTS sistema (
    chiave TEXT PRIMARY KEY,
    valore TEXT
)
""")

c.execute("""
CREATE TABLE IF NOT EXISTS sistema (
    chiave TEXT PRIMARY KEY,
    valore TEXT
)
""")

c.execute("""
CREATE TABLE IF NOT EXISTS chiusure (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    data TEXT NOT NULL,
    descrizione TEXT,
    corso_id INTEGER,
    attiva INTEGER DEFAULT 1
)
""")

c.execute("""
CREATE TABLE IF NOT EXISTS genitori_bambini (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    utente_id INTEGER NOT NULL,
    bambino_id INTEGER NOT NULL,
    UNIQUE(
        utente_id,
        bambino_id
    )
)
""")

c.execute("""
CREATE TABLE IF NOT EXISTS schede_genitori (

    id INTEGER PRIMARY KEY AUTOINCREMENT,

    bambino_id INTEGER UNIQUE,
    compilato INTEGER DEFAULT 0,
    bloccato INTEGER DEFAULT 0,
    data_invio TEXT,
    ultima_modifica TEXT,
    nome_genitore TEXT,
    cognome_genitore TEXT,
    luogo_nascita_genitore TEXT,
    data_nascita_genitore TEXT,
    cf_genitore TEXT,
    indirizzo_genitore TEXT,
    comune_genitore TEXT,
    cap_genitore TEXT
)
""")

try:

    c.execute("""
        ALTER TABLE schede_genitori
        ADD COLUMN data_nascita_bambino TEXT
    """)

    conn.commit()

except:
    pass

try:
    c.execute("""
        ALTER TABLE schede_genitori
        ADD COLUMN luogo_nascita_bambino TEXT
    """)
    conn.commit()
except:
    pass

try:
    c.execute("""
        ALTER TABLE schede_genitori
        ADD COLUMN cf_bambino TEXT
    """)
    conn.commit()
except:
    pass

try:
    c.execute("""
        ALTER TABLE schede_genitori
        ADD COLUMN indirizzo_bambino TEXT
    """)
    conn.commit()
except:
    pass

try:
    c.execute("""
        ALTER TABLE schede_genitori
        ADD COLUMN comune_bambino TEXT
    """)
    conn.commit()
except:
    pass

try:
    c.execute("""
        ALTER TABLE schede_genitori
        ADD COLUMN cap_bambino TEXT
    """)
    conn.commit()
except:
    pass

try:
    c.execute("""
        ALTER TABLE schede_genitori
        ADD COLUMN pratica_inviata INTEGER DEFAULT 0
    """)
except:
    pass

try:
    c.execute("""
        ALTER TABLE schede_genitori
        ADD COLUMN data_invio TEXT
    """)
except:
    pass

c.execute("""
CREATE TABLE IF NOT EXISTS storico_schede_genitore (

    id INTEGER PRIMARY KEY AUTOINCREMENT,

    bambino_id INTEGER,

    data_snapshot TEXT,

    json_snapshot TEXT
)
""")

conn.commit()
# ============================================================
# MIGRAZIONE CORSI VECCHI
# ============================================================

def invia_documentazione_email(
    bambino,
    modulo_firmato,
    certificato_medico,
    documento_identita,
    tessera_sanitaria,
    documento_identita_genitore
):

    msg = EmailMessage()
    
    msg["Subject"] = (
        f"Nuova iscrizione - "
        f"{bambino.iloc[0]['cognome']} "
        f"{bambino.iloc[0]['nome']} "
        f"- Stagione {stagione_corrente()}"
    )

    msg["From"] = st.secrets[
        "GMAIL_USER"
    ]

    msg["To"] = ", ".join(
        st.secrets["EMAIL_DESTINATARI"]
    )

    testo = f"""
        È stata inviata una nuova pratica di iscrizione relativa a: Corsi Power Team Piscina Comunale.
        
        =========================================
        DATI ATLETA
        =========================================
        
        Nome:
        {bambino.iloc[0]['nome']}
        
        Cognome:
        {bambino.iloc[0]['cognome']}
        
        Email genitore:
        {bambino.iloc[0]['email_genitore']}
        
        Telefono genitore:
        {bambino.iloc[0]['telefono_genitore']}
        
        Stagione:
        {stagione_corrente()}
        
        =========================================
        DOCUMENTI ALLEGATI
        =========================================
        
        - Modulo firmato
        
        - Documento identità bambino
        
        - Tessera sanitaria
        
        - Documento identità genitore
        
        {"- Certificato medico" if certificato_medico not in [None, True] else ""}
        
        =========================================
        
        Messaggio generato automaticamente dal gestionale Power Team Messina.
        """

    msg.set_content(
        testo
    )

    files = [
        ("modulo_firmato", modulo_firmato),
        ("certificato_medico", certificato_medico),
        ("documento_identita", documento_identita),
        ("tessera_sanitaria", tessera_sanitaria),
        ("documento_identita_genitore", documento_identita_genitore)
    ]

    for nome, file in files:

        if (
            file is None
            or file is True
        ):
            continue

        msg.add_attachment(
            file.getvalue(),
            maintype="application",
            subtype="octet-stream",
            filename=file.name
        )

    with smtplib.SMTP_SSL(
        "smtp.gmail.com",
        465
    ) as smtp:

        smtp.login(
            st.secrets["GMAIL_USER"],
            st.secrets["GMAIL_APP_PASSWORD"]
        )

        smtp.send_message(
            msg
        )
        
def crea_backup_completo():

    backup = {}

    tabelle = [
        "utenti",
        "corsi",
        "corso_giorni",
        "bambini",
        "genitori_bambini",
        "assegnazioni_istruttori",
        "presenze",
        "stagioni",
        "sistema",
        "chiusure"
    ]

    for tabella in tabelle:

        try:

            df = pd.read_sql(
                f"SELECT * FROM {tabella}",
                conn
            )

            backup[tabella] = df.to_dict(
                orient="records"
            )

        except Exception:

            backup[tabella] = []

    with open(
        "backup_completo.json",
        "w",
        encoding="utf-8"
    ) as f:

        json.dump(
            backup,
            f,
            ensure_ascii=False,
            indent=4
        )

    return "backup_completo.json"

def get_drive_service():
    
    service_account_info = {
        "type": "service_account",
        "project_id": st.secrets["GOOGLE_PROJECT_ID"],
        "private_key_id": st.secrets["GOOGLE_PRIVATE_KEY_ID"],
        "private_key": st.secrets["GOOGLE_PRIVATE_KEY"],
        "client_email": st.secrets["GOOGLE_CLIENT_EMAIL"],
        "client_id": st.secrets["GOOGLE_CLIENT_ID"],
        "token_uri": "https://oauth2.googleapis.com/token"
    }  
    
    credentials = service_account.Credentials.from_service_account_info(
        service_account_info,
        scopes=[
            "https://www.googleapis.com/auth/drive"
        ]
    )

    service = build(
        "drive",
        "v3",
        credentials=credentials
    )

    return service
    
def upload_backup_github(
    mostra_messaggio=True,
    forza=False
):

    global BACKUP_ABILITATO

    if not BACKUP_ABILITATO and not forza:
        return False

    token = st.secrets["GITHUB_TOKEN"]
    owner = st.secrets["GITHUB_OWNER"]
    repo = st.secrets["GITHUB_REPO"]

    path = st.secrets.get(
        "GITHUB_BACKUP_PATH",
        "backup_completo.json"
    )

    file_locale = crea_backup_completo()

    with open(
        file_locale,
        "r",
        encoding="utf-8"
    ) as f:

        contenuto = f.read()

    contenuto_b64 = base64.b64encode(
        contenuto.encode("utf-8")
    ).decode("utf-8")

    url = (
        f"https://api.github.com/repos/"
        f"{owner}/{repo}/contents/{path}"
    )

    headers = {
        "Authorization": f"Bearer {token}",
        "Accept": "application/vnd.github+json"
    }

    sha = None

    risposta = requests.get(
        url,
        headers=headers
    )
    
    if risposta.status_code == 200:
    
        sha = risposta.json().get("sha")

    payload = {
        "message": "Aggiornamento backup completo gestionale",
        "content": contenuto_b64
    }

    if sha is not None:

        payload["sha"] = sha
    
    upload = requests.put(
        url,
        headers=headers,
        json=payload
    )
    
    if upload.status_code == 409:
    
        risposta = requests.get(
            url,
            headers=headers
        )
    
        if risposta.status_code == 200:
    
            payload["sha"] = risposta.json()["sha"]
    
            upload = requests.put(
                url,
                headers=headers,
                json=payload
            )

    if upload.status_code in [200, 201]:

        if mostra_messaggio:

            st.success(
                "Backup salvato su GitHub."
            )

        return True

    else:

        if mostra_messaggio:

            st.error(
                f"Errore salvataggio GitHub: {upload.status_code}"
            )

            try:

                st.code(upload.text)

            except Exception:

                pass

        return False

def genera_pdf_presenze(
    bambino,
    corso,
    stagione,
    presenze_totali,
    assenze_totali,
    percentuale_presenza,
    data_generazione,
    df_calendario
):

    buffer = BytesIO()

    from reportlab.lib.pagesizes import landscape, A4

    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=15,
        rightMargin=15,
        topMargin=15,
        bottomMargin=15
    )

    styles = getSampleStyleSheet()

    elementi = []

    elementi.append(
        Paragraph(
            "POWER TEAM MESSINA",
            styles["Title"]
        )
    )

    elementi.append(
        Paragraph(
            "Riepilogo Presenze Stagionale",
            styles["Heading2"]
        )
    )

    elementi.append(Spacer(1,12))

    elementi.append(
        Paragraph(
            f"""
            <b>Bambino:</b> {bambino}<br/>
            <b>Corso:</b> {corso}<br/>
            <b>Stagione:</b> {stagione}<br/>
            <b>Presenze:</b> {presenze_totali}<br/>
            <b>Assenze:</b> {assenze_totali}<br/>
            <b>Percentuale presenza:</b> {percentuale_presenza}%<br/>
            <b>Data generazione:</b> {data_generazione}
            """,
            styles["BodyText"]
        )
    )

    elementi.append(Spacer(1,12))

    df_pdf = df_calendario.copy()

    df_pdf = df_pdf.replace(
        {
            "✅": "P",
            "❌": "A"
        }
    )
    
    dati_tabella = [
        df_pdf.columns.tolist()
    ]
    
    dati_tabella += (
        df_pdf.values.tolist()
    )

    tabella = Table(
        dati_tabella
    )

    tabella.setStyle(
        TableStyle([
            ("GRID",(0,0),(-1,-1),0.5,colors.black),
            ("BACKGROUND",(0,0),(-1,0),colors.lightgrey),
            ("FONTSIZE",(0,0),(-1,-1),7)
        ])
    )

    elementi.append(tabella)

    elementi.append(Spacer(1,12))

    elementi.append(
        Paragraph(
            "Legenda: P = Presenza | A = Assenza | - = Nessuna lezione / Festività / Chiusura",
            styles["BodyText"]
        )
    )

    doc.build(elementi)

    pdf = buffer.getvalue()

    buffer.close()

    return pdf

def pratica_completa(
    scheda,
    modulo_firmato,
    certificato_medico,
    documento_identita,
    tessera_sanitaria,
    documento_identita_genitore
):

    if scheda.empty:
        return False

    s = scheda.iloc[0]

    campi = [
        "nome_genitore",
        "cognome_genitore",
        "cf_genitore",
        "indirizzo_genitore",
        "comune_genitore",
        "cap_genitore"
    ]

    for campo in campi:

        if (
            pd.isna(
                s[campo]
            )
            or
            str(
                s[campo]
            ).strip() == ""
        ):
            return False

    if modulo_firmato is None:
        return False

    if certificato_medico is None:
        return False

    if documento_identita is None:
        return False

    if tessera_sanitaria is None:
        return False

    if documento_identita_genitore is None:
        return False

    return True

def scarica_file_drive_bytes(
    file_id
):

    service = get_drive_service()

    request = service.files().get_media(
        fileId=file_id
    )

    buffer = io.BytesIO()

    downloader = MediaIoBaseDownload(
        buffer,
        request
    )

    done = False

    while not done:

        _, done = downloader.next_chunk()

    buffer.seek(0)

    return buffer

def valore_pulito(
    valore
):

    if pd.isna(
        valore
    ):

        return ""

    return str(
        valore
    )

def formatta_data_modulo(
    valore
):

    if pd.isna(
        valore
    ):

        return ""

    try:

        return pd.to_datetime(
            valore,
            errors="coerce"
        ).strftime(
            "%d/%m/%Y"
        )

    except Exception:

        return str(
            valore
        )

def sostituisci_testo_paragrafo(
    paragrafo,
    mappa
):

    testo_originale = paragrafo.text

    testo_nuovo = testo_originale

    for chiave, valore in mappa.items():

        testo_nuovo = testo_nuovo.replace(
            chiave,
            str(
                valore
            )
        )

    if testo_nuovo != testo_originale:

        for run in paragrafo.runs:

            run.text = ""

        if paragrafo.runs:

            paragrafo.runs[0].text = testo_nuovo

        else:

            paragrafo.add_run(
                testo_nuovo
            )

def compila_docx_template(
    template_bytes,
    mappa
):

    with tempfile.NamedTemporaryFile(
        suffix=".docx",
        delete=False
    ) as tmp:

        tmp.write(
            template_bytes.getvalue()
        )

        tmp_path = tmp.name

    doc = Document(
        tmp_path
    )

    for paragrafo in doc.paragraphs:

        sostituisci_testo_paragrafo(
            paragrafo,
            mappa
        )

    for tabella in doc.tables:

        for riga in tabella.rows:

            for cella in riga.cells:

                for paragrafo in cella.paragraphs:

                    sostituisci_testo_paragrafo(
                        paragrafo,
                        mappa
                    )

    output = io.BytesIO()

    doc.save(
        output
    )

    output.seek(0)

    return output

def scrivi_centrato(
    c,
    testo,
    centro_x,
    y,
    font="Helvetica",
    size=8
):

    c.setFont(
        font,
        size
    )

    larghezza = c.stringWidth(
        str(testo),
        font,
        size
    )

    c.drawString(
        centro_x - larghezza / 2,
        y,
        str(testo)
    )
    
def converti_docx_in_pdf_drive(
    docx_bytes,
    nome_file="modulo_compilato.docx"
):

    service = get_drive_service()

    docx_bytes.seek(0)

    media = MediaIoBaseUpload(
        docx_bytes,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        resumable=False
    )

    metadata = {
        "name": nome_file,
        "mimeType": "application/vnd.google-apps.document",
        "parents": [
            st.secrets["DRIVE_FOLDER_ID"]
        ]
    }

    st.info("Sto creando il Google Document temporaneo...")

    file_creato = service.files().create(
        body=metadata,
        media_body=media,
        fields="id"
    ).execute()

    google_doc_id = file_creato["id"]

    request = service.files().export_media(
        fileId=google_doc_id,
        mimeType="application/pdf"
    )

    pdf_buffer = io.BytesIO()

    downloader = MediaIoBaseDownload(
        pdf_buffer,
        request
    )

    done = False

    while not done:

        _, done = downloader.next_chunk()

    pdf_buffer.seek(0)

    try:

        service.files().delete(
            fileId=google_doc_id
        ).execute()

    except Exception:

        pass

    return pdf_buffer

def crea_mappa_modulo(
    bambino,
    scheda
):

    b = bambino.iloc[0]
    s = scheda.iloc[0]

    mappa = {

        "{{NOME_BAMBINO}}": valore_pulito(
            b["nome"]
        ),

        "{{COGNOME_BAMBINO}}": valore_pulito(
            b["cognome"]
        ),

        "{{DATA_NASCITA_BAMBINO}}": formatta_data_modulo(
            s.get(
                "data_nascita_bambino",
                ""
            )
        ),

        "{{LUOGO_NASCITA_BAMBINO}}": valore_pulito(
            s.get(
                "luogo_nascita_bambino",
                ""
            )
        ),

        "{{CF_BAMBINO}}": valore_pulito(
            s.get(
                "cf_bambino",
                ""
            )
        ),

        "{{INDIRIZZO_BAMBINO}}": valore_pulito(
            s.get(
                "indirizzo_bambino",
                ""
            )
        ),

        "{{COMUNE_RESIDENZA_BAMBINO}}": valore_pulito(
            s.get("comune_bambino", "")
        ),
        
        "{{COMUNE_RESIDENZA_GENITORE}}": valore_pulito(
            s.get("comune_genitore", "")
        ),
        
        "{{EMAIL_RIFERIMENTO}}": valore_pulito(
            b.get("email_genitore", "")
        ),
        
        "{{TELEFONO_RIFERIMENTO}}": valore_pulito(
            b.get("telefono_genitore", "")
        ),
        
        "{{EMAIL_GENITORE}}": valore_pulito(
            b.get("email_genitore", "")
        ),
        
        "{{TELEFONO_GENITORE}}": valore_pulito(
            b.get("telefono_genitore", "")
        ),
        
        "{{CAP_BAMBINO}}": valore_pulito(
            s.get(
                "cap_bambino",
                ""
            )
        ),

        "{{NOME_GENITORE}}": valore_pulito(
            s.get(
                "nome_genitore",
                ""
            )
        ),

        "{{COGNOME_GENITORE}}": valore_pulito(
            s.get(
                "cognome_genitore",
                ""
            )
        ),

        "{{LUOGO_NASCITA_GENITORE}}": valore_pulito(
            s.get(
                "luogo_nascita_genitore",
                ""
            )
        ),

        "{{DATA_NASCITA_GENITORE}}": formatta_data_modulo(
            s.get(
                "data_nascita_genitore",
                ""
            )
        ),

        "{{CF_GENITORE}}": valore_pulito(
            s.get(
                "cf_genitore",
                ""
            )
        ),

        "{{INDIRIZZO_GENITORE}}": valore_pulito(
            s.get(
                "indirizzo_genitore",
                ""
            )
        ),

        "{{CAP_GENITORE}}": valore_pulito(
            s.get(
                "cap_genitore",
                ""
            )
        ),

        "{{DATA_MODULO}}": datetime.now().strftime(
            "%d/%m/%Y"
        )
    }

    return mappa

def get_info_corso(
    corso_id
):

    return pd.read_sql(
        """
        SELECT
            nome,
            giorno,
            orario,
            stagione
        FROM corsi
        WHERE id = ?
        """,
        conn,
        params=(corso_id,)
    )
    
def scarica_backup_github():

    token = st.secrets["GITHUB_TOKEN"]
    owner = st.secrets["GITHUB_OWNER"]
    repo = st.secrets["GITHUB_REPO"]

    path = st.secrets.get(
        "GITHUB_BACKUP_PATH",
        "backup_completo.json"
    )

    url = (
        f"https://api.github.com/repos/"
        f"{owner}/{repo}/contents/{path}"
    )

    headers = {
        "Authorization": f"Bearer {token}",
        "Accept": "application/vnd.github+json"
    }

    risposta = requests.get(
        url,
        headers=headers
    )

    if risposta.status_code != 200:

        return False

    data = risposta.json()

    contenuto = base64.b64decode(
        data["content"]
    ).decode("utf-8")

    with open(
        "backup_completo.json",
        "w",
        encoding="utf-8"
    ) as f:

        f.write(contenuto)

    return True

def is_festivo(data):

    festivita = {

        "01-01",
        "06-01",

        "25-04",

        "01-05",

        "02-06",

        "15-08",

        "01-11",

        "08-12",

        "25-12",
        "26-12"
    }

    return (
        data.strftime("%m-%d")
        in festivita
    )


def ripristina_backup_locale():

    if not os.path.exists(
        "backup_completo.json"
    ):

        return False

    with open(
        "backup_completo.json",
        "r",
        encoding="utf-8"
    ) as f:

        backup = json.load(f)

    ordine_delete = [
        "presenze",
        "assegnazioni_istruttori",
        "bambini",
        "genitori_bambini",
        "corso_giorni",
        "corsi",
        "utenti",
        "stagioni",
        "sistema",
        "chiusure"
    ]

    for tabella in ordine_delete:

        try:

            c.execute(
                f"DELETE FROM {tabella}"
            )

        except Exception as e:

            print(
                f"Errore cancellazione tabella {tabella}: {e}"
            )

    conn.commit()

    ordine_insert = [
        "utenti",
        "stagioni",
        "corsi",
        "corso_giorni",
        "bambini",
        "genitori_bambini",
        "assegnazioni_istruttori",
        "presenze",
        "sistema",
        "chiusure"
    ]

    for tabella in ordine_insert:

        records = backup.get(
            tabella,
            []
        )

        if len(records) == 0:

            continue

        df = pd.DataFrame(records)

        try:

            colonne_db = pd.read_sql(
                f"PRAGMA table_info({tabella})",
                conn
            )["name"].tolist()

            colonne_valide = [
                col
                for col in df.columns
                if col in colonne_db
            ]

            df = df[
                colonne_valide
            ]

            if df.empty:

                continue

            df.to_sql(
                tabella,
                conn,
                if_exists="append",
                index=False
            )

        except Exception as e:

            print(
                f"ERRORE RIPRISTINO TABELLA {tabella}: {e}"
            )

            raise

    conn.commit()

    return True

def database_vuoto():

    try:

        utenti = pd.read_sql(
            """
            SELECT COUNT(*) AS totale
            FROM utenti
            """,
            conn
        ).iloc[0]["totale"]

        corsi = pd.read_sql(
            """
            SELECT COUNT(*) AS totale
            FROM corsi
            """,
            conn
        ).iloc[0]["totale"]

        bambini = pd.read_sql(
            """
            SELECT COUNT(*) AS totale
            FROM bambini
            """,
            conn
        ).iloc[0]["totale"]

        return (
            utenti == 0
            and corsi == 0
            and bambini == 0
        )

    except Exception:

        return True

def ripristino_iniziale_da_github():

    if database_vuoto():
    
        try:
    
            if scarica_backup_github():
    
                ripristina_backup_locale()
    
        except Exception as e:
    
            print(
                f"Errore ripristino backup: {e}"
            )

def backup_giornaliero_github():

    oggi = datetime.today().strftime(
        "%Y-%m-%d"
    )

    df = pd.read_sql(
        """
        SELECT valore
        FROM sistema
        WHERE chiave = 'ultimo_backup_github'
        """,
        conn
    )

    if (
        df.empty
        or df.iloc[0]["valore"] != oggi
    ):

        ok = upload_backup_github(
            mostra_messaggio=False
        )

        if ok:

            c.execute(
                """
                INSERT OR REPLACE INTO sistema(
                    chiave,
                    valore
                )
                VALUES(?,?)
                """,
                (
                    "ultimo_backup_github",
                    oggi
                )
            )

            conn.commit()

            try:
        
                upload_backup_github(
                    mostra_messaggio=False
                )
            
            except Exception as e:
            
                print(
                    f"Errore backup GitHub: {e}"
                )

def migra_giorni_corsi_vecchi():

    corsi_vecchi = pd.read_sql(
        """
        SELECT id, giorno, orario
        FROM corsi
        WHERE giorno IS NOT NULL
        AND giorno != ''
        """,
        conn
    )

    for _, row in corsi_vecchi.iterrows():

        esiste = pd.read_sql(
            """
            SELECT *
            FROM corso_giorni
            WHERE corso_id = ?
            """,
            conn,
            params=(int(row["id"]),)
        )

        if esiste.empty:

            c.execute(
                """
                INSERT INTO corso_giorni(
                    corso_id,
                    giorno,
                    orario,
                    ordine
                )
                VALUES(?,?,?,1)
                """,
                (
                    int(row["id"]),
                    row["giorno"],
                    row["orario"] if pd.notna(row["orario"]) else ""
                )
            )

    conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )
    
    except Exception as e:
    
        print(
            f"Errore backup GitHub: {e}"
        )

#migra_giorni_corsi_vecchi()

# ============================================================
# CREAZIONE MANAGER DEFAULT
# ============================================================

def crea_manager_default():

    df = pd.read_sql(
        """
        SELECT *
        FROM utenti
        WHERE ruolo = 'manager'
        """,
        conn
    )

    if df.empty:

        username = st.secrets.get(
            "MANAGER_USERNAME",
            "manager"
        )

        password = st.secrets.get(
            "MANAGER_PASSWORD",
            "manager123"
        )

        password_hash, salt = hash_password(password)

        c.execute(
            """
            INSERT INTO utenti(
                username,
                nome,
                ruolo,
                password_hash,
                salt,
                attivo
            )
            VALUES(?,?,?,?,?,1)
            """,
            (
                username,
                "Manager",
                "manager",
                password_hash,
                salt
            )
        )

        conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )
    
    except Exception as e:
    
        print(
            f"Errore backup GitHub: {e}"
        )

#crea_manager_default()

def get_bambini_corso(corso_id, attivi_solo=True):

    query = """
        SELECT *
        FROM bambini
        WHERE corso_id = ?
    """

    if attivi_solo:
        query += " AND attivo = 1"

    query += " ORDER BY cognome, nome"

    try:

        return pd.read_sql(
            query,
            conn,
            params=(corso_id,)
        )

    except Exception as e:

        st.error(str(e))
        raise

def get_bambini(attivi_solo=True):

    query = """
        SELECT *
        FROM bambini
        WHERE 1 = 1
    """

    if attivi_solo:
        query += """
            AND attivo = 1
        """

    query += """
        ORDER BY cognome, nome
    """

    return pd.read_sql(
        query,
        conn
    )
    
def aggiungi_stagione(nome):

    c.execute(
        """
        INSERT INTO stagioni(
            nome,
            attiva
        )
        VALUES(?,1)
        """,
        (nome,)
    )

    conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )
    
    except Exception as e:
    
        print(
            f"Errore backup GitHub: {e}"
        )

def get_riepilogo_stagioni():

    # --------------------------------------------------------
    # Recupero stagioni
    # --------------------------------------------------------

    try:

        stagioni_df = pd.read_sql(
            """
            SELECT nome
            FROM stagioni
            WHERE attiva = 1
            ORDER BY nome DESC
            """,
            conn
        )

    except Exception:

        stagioni_df = pd.read_sql(
            """
            SELECT DISTINCT stagione AS nome
            FROM corsi
            WHERE stagione IS NOT NULL
            AND stagione != ''
            ORDER BY stagione DESC
            """,
            conn
        )

    risultati = []

    for _, stagione_row in stagioni_df.iterrows():

        stagione = stagione_row["nome"]

        # ----------------------------------------------------
        # Numero corsi
        # ----------------------------------------------------

        numero_corsi = pd.read_sql(
            """
            SELECT COUNT(*) AS totale
            FROM corsi
            WHERE stagione = ?
            """,
            conn,
            params=(stagione,)
        ).iloc[0]["totale"]

        # ----------------------------------------------------
        # Numero istruttori assegnati ai corsi della stagione
        # ----------------------------------------------------

        numero_istruttori = pd.read_sql(
            """
            SELECT COUNT(DISTINCT ai.istruttore_id) AS totale
            FROM assegnazioni_istruttori ai
            JOIN corsi c
                ON c.id = ai.corso_id
            WHERE c.stagione = ?
            AND ai.attiva = 1
            """,
            conn,
            params=(stagione,)
        ).iloc[0]["totale"]

        # ----------------------------------------------------
        # Numero bambini
        # ----------------------------------------------------
        # Proviamo prima con corso_id nella tabella bambini.
        # Se la colonna non esiste, usiamo assegnazioni_bambini.
        # ----------------------------------------------------

        colonne_bambini = pd.read_sql(
            """
            PRAGMA table_info(bambini)
            """,
            conn
        )["name"].tolist()

        if "corso_id" in colonne_bambini:

            numero_bambini = pd.read_sql(
                """
                SELECT COUNT(DISTINCT b.id) AS totale
                FROM bambini b
                JOIN corsi c
                    ON c.id = b.corso_id
                WHERE c.stagione = ?
                AND b.attivo = 1
                """,
                conn,
                params=(stagione,)
            ).iloc[0]["totale"]

        else:

            try:

                numero_bambini = pd.read_sql(
                    """
                    SELECT COUNT(DISTINCT ab.bambino_id) AS totale
                    FROM assegnazioni_bambini ab
                    JOIN corsi c
                        ON c.id = ab.corso_id
                    JOIN bambini b
                        ON b.id = ab.bambino_id
                    WHERE c.stagione = ?
                    AND ab.attiva = 1
                    AND b.attivo = 1
                    """,
                    conn,
                    params=(stagione,)
                ).iloc[0]["totale"]

            except Exception:

                numero_bambini = 0

        risultati.append(
            {
                "stagione": stagione,
                "numero_corsi": int(numero_corsi),
                "numero_istruttori": int(numero_istruttori),
                "numero_bambini": int(numero_bambini)
            }
        )

    return pd.DataFrame(
        risultati
    )

def elimina_presenze_giornata(
    corso_id,
    data_evento
):

    c.execute(
        """
        DELETE FROM presenze
        WHERE corso_id = ?
        AND data = ?
        """,
        (
            corso_id,
            str(data_evento)
        )
    )

    numero = c.rowcount

    conn.commit()

    return numero
# ============================================================
# FUNZIONI UTENTI
# ============================================================
def invia_credenziali_istruttore_email(
    email_destinatario,
    nome,
    password
):

    mittente = st.secrets["GMAIL_ADDRESS"]
    password_mittente = st.secrets["GMAIL_PASSWORD"]

    app_url = st.secrets.get(
        "APP_URL",
        "link dell'app Streamlit"
    )

    testo = f"""
Ciao {nome},

il tuo account istruttore per lo Statino Corsi Nuoto è stato creato.

Credenziali di accesso:

Email: {email_destinatario}
Password: {password}

Link app:
{app_url}

Ti consigliamo di conservare queste credenziali con attenzione.

Power Team Messina
"""

    msg = MIMEText(
        testo,
        "plain",
        "utf-8"
    )

    msg["Subject"] = "Credenziali accesso Statino Corsi Nuoto"
    msg["From"] = mittente
    msg["To"] = email_destinatario

    server = smtplib.SMTP(
        "smtp.gmail.com",
        587
    )

    server.starttls()

    server.login(
        mittente,
        password_mittente
    )

    server.send_message(msg)

    server.quit()

def invia_credenziali_manager_email(
    email_destinatario,
    nome,
    password
):

    mittente = st.secrets["GMAIL_ADDRESS"]
    password_mittente = st.secrets["GMAIL_PASSWORD"]

    app_url = st.secrets.get(
        "APP_URL",
        ""
    )

    testo = f"""
Ciao {nome},

è stato creato un nuovo account MANAGER per l'App Corsi Power Team.

Credenziali di accesso:

Email: {email_destinatario}
Password: {password}

Link applicazione:
{app_url}

Con questo account avrai accesso completo alla piattaforma.

Power Team Messina
"""

    msg = MIMEText(
        testo,
        "plain",
        "utf-8"
    )

    msg["Subject"] = "Credenziali Manager Gestionale Nuoto"
    msg["From"] = mittente
    msg["To"] = email_destinatario

    server = smtplib.SMTP(
        "smtp.gmail.com",
        587
    )

    server.starttls()

    server.login(
        mittente,
        password_mittente
    )

    server.send_message(msg)

    server.quit()
    
def get_utenti():

    return pd.read_sql(
        """
        SELECT
            id,
            username,
            nome,
            ruolo,
            attivo
        FROM utenti
        ORDER BY ruolo, nome
        """,
        conn
    )


def get_istruttori(attivi_solo=True):

    query = """
        SELECT
            id,
            username,
            nome,
            ruolo,
            attivo
        FROM utenti
        WHERE ruolo = 'istruttore'
    """

    if attivi_solo:
        query += " AND attivo = 1"

    query += " ORDER BY nome"

    return pd.read_sql(
        query,
        conn
    )

def aggiungi_istruttore(email, nome):

    password_generata = genera_password_casuale()

    password_hash, salt = hash_password(
        password_generata
    )

    c.execute(
        """
        INSERT INTO utenti(
            username,
            nome,
            ruolo,
            password_hash,
            salt,
            password_visibile,
            attivo
        )
        VALUES(?,?,?,?,?,?,1)
        """,
        (
            email.strip().lower(),
            nome.strip(),
            "istruttore",
            password_hash,
            salt,
            password_generata
        )
    )

    conn.commit()

    try:

        ok = upload_backup_github(
            mostra_messaggio=True
        )
    
        st.write(
            f"Backup GitHub: {ok}"
        )
    
    except Exception as e:
    
        st.error(
            f"Errore backup GitHub: {e}"
        )

    return password_generata

def aggiungi_manager(email, nome):

    password_generata = genera_password_casuale()

    password_hash, salt = hash_password(
        password_generata
    )

    c.execute(
        """
        INSERT INTO utenti(
            username,
            nome,
            ruolo,
            password_hash,
            salt,
            password_visibile,
            attivo
        )
        VALUES(?,?,?,?,?,?,1)
        """,
        (
            email.strip().lower(),
            nome.strip(),
            "manager",
            password_hash,
            salt,
            password_generata
        )
    )

    conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )

    except:
        pass

    return password_generata

def aggiorna_password_utente(
    utente_id,
    nuova_password
):

    password_hash, salt = hash_password(
        nuova_password
    )

    c.execute(
        """
        UPDATE utenti
        SET
            password_hash = ?,
            salt = ?,
            password_visibile = ?
        WHERE id = ?
        """,
        (
            password_hash,
            salt,
            nuova_password,
            utente_id
        )
    )

    conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )
    
    except Exception as e:
    
        print(
            f"Errore backup GitHub: {e}"
        )

def cambia_stato_utente(utente_id, attivo):

    c.execute(
        """
        UPDATE utenti
        SET attivo = ?
        WHERE id = ?
        """,
        (
            int(attivo),
            utente_id
        )
    )

    conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )
    
    except Exception as e:
    
        print(
            f"Errore backup GitHub: {e}"
        )

def login():

    st.sidebar.header("🔐 Login")

    if "utente_id" not in st.session_state:
        st.session_state.utente_id = None

    if "username" not in st.session_state:
        st.session_state.username = None

    if "nome_utente" not in st.session_state:
        st.session_state.nome_utente = None

    if "ruolo" not in st.session_state:
        st.session_state.ruolo = None

    if st.session_state.utente_id is not None:

        st.sidebar.success(
            f"Accesso: {st.session_state.nome_utente}"
        )

        st.sidebar.write(
            f"Ruolo: {st.session_state.ruolo}"
        )

        if st.sidebar.button("🚪 Logout"):

            st.session_state.utente_id = None
            st.session_state.username = None
            st.session_state.nome_utente = None
            st.session_state.ruolo = None

            st.rerun()

        return

    username = st.sidebar.text_input(
        "Email / username"
    )
    
    password = st.sidebar.text_input(
        "Password",
        type="password"
    )
    
    if st.sidebar.button("Accedi"):
    
        username_login = username.strip().lower()
    
        utente = pd.read_sql(
            """
            SELECT *
            FROM utenti
            WHERE lower(username) = ?
            AND attivo = 1
            """,
            conn,
            params=(username.strip().lower(),)
        )

        if utente.empty:

            st.sidebar.error(
                "Credenziali non valide."
            )

            return

        row = utente.iloc[0]

        ok = verifica_password(
            password,
            row["password_hash"],
            row["salt"]
        )

        if ok:

            st.session_state.utente_id = int(row["id"])
            st.session_state.username = row["username"]
            st.session_state.nome_utente = row["nome"]
            st.session_state.ruolo = row["ruolo"]

            st.rerun()

        else:

            st.sidebar.error(
                "Credenziali non valide."
            )

def is_genitore():

    return (
        st.session_state.get(
            "ruolo"
        ) == "genitore"
    )

def is_manager():

    return st.session_state.get("ruolo") == "manager"


def is_istruttore():

    return st.session_state.get("ruolo") == "istruttore"


def is_loggato():

    return st.session_state.get("utente_id") is not None


# ============================================================
# FUNZIONI CORSI
# ============================================================

def get_corsi(
    attivi_solo=True,
    stagione=None
):

    query = """
        SELECT *
        FROM corsi
        WHERE 1=1
    """

    params = []

    if attivi_solo:

        query += """
            AND attivo = 1
        """

    if stagione is not None:

        query += """
            AND stagione = ?
        """

        params.append(stagione)

    query += """
        ORDER BY nome
    """

    return pd.read_sql(
        query,
        conn,
        params=params
    )

def get_giorni_corso(corso_id):

    return pd.read_sql(
        """
        SELECT *
        FROM corso_giorni
        WHERE corso_id = ?
        ORDER BY ordine, giorno, orario
        """,
        conn,
        params=(corso_id,)
    )

def descrizione_giorni_corso(corso_id):

    giorni = get_giorni_corso(corso_id)

    if giorni.empty:
        return "Nessun giorno assegnato"

    testi = []

    for _, row in giorni.iterrows():

        testi.append(
            f"{row['giorno']} {row['orario']}"
        )

    return " | ".join(testi)


def get_corsi_con_giorni(
    attivi_solo=True,
    stagione=None
):

    corsi = get_corsi(
        attivi_solo=attivi_solo,
        stagione=stagione
    )

    if corsi.empty:
        return corsi

    descrizioni = []

    for _, row in corsi.iterrows():

        descrizioni.append(
            descrizione_giorni_corso(
                int(row["id"])
            )
        )

    corsi["giorni_orari"] = descrizioni

    return corsi

def get_corso_by_id(corso_id):

    return pd.read_sql(
        """
        SELECT *
        FROM corsi
        WHERE id = ?
        """,
        conn,
        params=(corso_id,)
    )
        
def aggiungi_corso(nome, livello, stagione):

    c.execute(
        """
        INSERT INTO corsi(
            nome,
            livello,
            giorno,
            orario,
            stagione,
            attivo
        )
        VALUES(?,?,?,?,?,1)
        """,
        (
            nome,
            livello,
            "",
            "",
            stagione
        )
    )

    conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )
    
    except Exception as e:
    
        print(
            f"Errore backup GitHub: {e}"
        )
        
    return c.lastrowid

def salva_giorni_corso(corso_id, giorni_orari):

    c.execute(
        """
        DELETE FROM corso_giorni
        WHERE corso_id = ?
        """,
        (corso_id,)
    )

    ordine = 1

    for giorno, orario in giorni_orari:

        if giorno.strip() != "" and orario.strip() != "":

            c.execute(
                """
                INSERT INTO corso_giorni(
                    corso_id,
                    giorno,
                    orario,
                    ordine
                )
                VALUES(?,?,?,?)
                """,
                (
                    corso_id,
                    giorno.strip(),
                    orario.strip(),
                    ordine
                )
            )

            ordine += 1

    conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )
    
    except Exception as e:
    
        print(
            f"Errore backup GitHub: {e}"
        )

def crea_account_genitore(
    email,
    nome_bambino
):

    email = email.strip().lower()

    esistente = pd.read_sql(
        """
        SELECT *
        FROM utenti
        WHERE lower(username)=?
        """,
        conn,
        params=(email,)
    )

    if not esistente.empty:

        return (
            int(esistente.iloc[0]["id"]),
            esistente.iloc[0]["password_visibile"],
            True
        )

    password_generata = genera_password_casuale()

    password_hash, salt = hash_password(
        password_generata
    )

    c.execute(
        """
        INSERT INTO utenti(
            username,
            nome,
            ruolo,
            password_hash,
            salt,
            password_visibile,
            attivo
        )
        VALUES(?,?,?,?,?,?,1)
        """,
        (
            email,
            f"Genitore {nome_bambino}",
            "genitore",
            password_hash,
            salt,
            password_generata
        )
    )

    utente_id = c.lastrowid

    conn.commit()

    return (
        utente_id,
        password_generata,
        False
    )

def aggiorna_corso(corso_id, nome, livello, stagione, attivo):

    c.execute(
        """
        UPDATE corsi
        SET nome = ?,
            livello = ?,
            stagione = ?,
            attivo = ?
        WHERE id = ?
        """,
        (
            nome,
            livello,
            stagione,
            int(attivo),
            corso_id
        )
    )

    conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )
    
    except Exception as e:
    
        print(
            f"Errore backup GitHub: {e}"
        )

def elimina_corso(corso_id):

    c.execute(
        """
        DELETE FROM presenze
        WHERE corso_id = ?
        """,
        (corso_id,)
    )

    c.execute(
        """
        DELETE FROM bambini
        WHERE corso_id = ?
        """,
        (corso_id,)
    )

    c.execute(
        """
        DELETE FROM assegnazioni_istruttori
        WHERE corso_id = ?
        """,
        (corso_id,)
    )

    c.execute(
        """
        DELETE FROM corso_giorni
        WHERE corso_id = ?
        """,
        (corso_id,)
    )

    c.execute(
        """
        DELETE FROM corsi
        WHERE id = ?
        """,
        (corso_id,)
    )

    conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )
    
    except Exception as e:
    
        print(
            f"Errore backup GitHub: {e}"
        )

def giorno_settimana_italiano(data_evento):

    giorni = {
        0: "Lunedì",
        1: "Martedì",
        2: "Mercoledì",
        3: "Giovedì",
        4: "Venerdì",
        5: "Sabato",
        6: "Domenica"
    }

    return giorni[pd.Timestamp(data_evento).weekday()]

def genera_pdf_riepilogo_mensile_corso(
    nome_corso,
    mese_nome,
    anno,
    stagione,
    df_mensile,
    data_generazione
):

    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=15,
        rightMargin=15,
        topMargin=15,
        bottomMargin=15
    )

    styles = getSampleStyleSheet()

    elementi = []

    elementi.append(
        Paragraph(
            "POWER TEAM MESSINA",
            styles["Title"]
        )
    )

    elementi.append(
        Paragraph(
            "Riepilogo Presenze Mensile per Corso",
            styles["Heading2"]
        )
    )

    elementi.append(
        Spacer(
            1,
            12
        )
    )

    elementi.append(
        Paragraph(
            f"""
            <b>Corso:</b> {nome_corso}<br/>
            <b>Mese:</b> {mese_nome} {anno}<br/>
            <b>Stagione:</b> {stagione}<br/>
            <b>Data generazione:</b> {data_generazione}
            """,
            styles["BodyText"]
        )
    )

    elementi.append(
        Spacer(
            1,
            12
        )
    )

    dati_tabella = [
        df_mensile.columns.tolist()
    ]

    dati_tabella += df_mensile.values.tolist()

    tabella = Table(
        dati_tabella,
        repeatRows=1
    )

    stile = TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.black),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 6),
        ("ALIGN", (1, 1), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE")
    ])

    for riga_idx, riga in enumerate(
        dati_tabella[1:],
        start=1
    ):

        for col_idx, valore in enumerate(
            riga
        ):

            if valore == "P":

                stile.add(
                    "BACKGROUND",
                    (col_idx, riga_idx),
                    (col_idx, riga_idx),
                    colors.lightgreen
                )

            elif valore == "A":

                stile.add(
                    "BACKGROUND",
                    (col_idx, riga_idx),
                    (col_idx, riga_idx),
                    colors.lightcoral
                )

            elif valore == "-":

                stile.add(
                    "BACKGROUND",
                    (col_idx, riga_idx),
                    (col_idx, riga_idx),
                    colors.whitesmoke
                )

    tabella.setStyle(
        stile
    )

    elementi.append(
        tabella
    )

    elementi.append(
        Spacer(
            1,
            12
        )
    )

    elementi.append(
        Paragraph(
            "Legenda: P = Presenza | A = Assenza | - = Nessuna lezione / Festività / Chiusura",
            styles["BodyText"]
        )
    )

    doc.build(
        elementi
    )

    pdf = buffer.getvalue()

    buffer.close()

    return pdf
    
# ============================================================
# FUNZIONI ASSEGNAZIONI
# ============================================================

def assegna_istruttore(istruttore_id, corso_id, data_specifica=None):

    if data_specifica == "":
        data_specifica = None

    c.execute(
        """
        INSERT INTO assegnazioni_istruttori(
            istruttore_id,
            corso_id,
            data_specifica,
            attiva
        )
        VALUES(?,?,?,1)
        """,
        (
            istruttore_id,
            corso_id,
            data_specifica
        )
    )

    conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )
    
    except Exception as e:
    
        print(
            f"Errore backup GitHub: {e}"
        )


def get_assegnazioni():

    assegnazioni = pd.read_sql(
        """
        SELECT
            ai.id,
            u.nome AS istruttore,
            c.id AS corso_id,
            c.nome AS corso,
            ai.data_specifica,
            ai.attiva
        FROM assegnazioni_istruttori ai
        JOIN utenti u
            ON u.id = ai.istruttore_id
        JOIN corsi c
            ON c.id = ai.corso_id
        ORDER BY u.nome, c.nome
        """,
        conn
    )

    giorni_descrizione = []

    for _, row in assegnazioni.iterrows():

        giorni = get_giorni_corso(
            int(row["corso_id"])
        )

        if giorni.empty:

            giorni_descrizione.append(
                ""
            )

        else:

            giorni_descrizione.append(
                " | ".join(
                    [
                        f"{g['giorno']} {g['orario']}"
                        for _, g in giorni.iterrows()
                    ]
                )
            )

    assegnazioni["giorni"] = (
        giorni_descrizione
    )

    return assegnazioni


def elimina_assegnazione(assegnazione_id):

    c.execute(
        """
        DELETE FROM assegnazioni_istruttori
        WHERE id = ?
        """,
        (assegnazione_id,)
    )

    conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )
    
    except Exception as e:
    
        print(
            f"Errore backup GitHub: {e}"
        )


def istruttore_abilitato_corso_data(istruttore_id, corso_id, data_evento):

    data_str = str(data_evento)

    df = pd.read_sql(
        """
        SELECT *
        FROM assegnazioni_istruttori
        WHERE istruttore_id = ?
        AND corso_id = ?
        AND attiva = 1
        AND (
            data_specifica IS NULL
            OR data_specifica = ?
        )
        """,
        conn,
        params=(
            istruttore_id,
            corso_id,
            data_str
        )
    )

    return not df.empty


def get_corsi_visibili_per_utente(data_evento=None):

    filtro_giorno = ""

    params_extra = []

    if data_evento is not None:

        giorno_evento = giorno_settimana_italiano(
            data_evento
        )

        filtro_giorno = """
            AND cg.giorno = ?
        """

        params_extra.append(giorno_evento)

    if is_manager():

        query = f"""
            SELECT DISTINCT
                c.*,
                cg.giorno AS giorno_lezione,
                cg.orario AS orario_lezione
            FROM corsi c
            JOIN corso_giorni cg
                ON cg.corso_id = c.id
            WHERE c.attivo = 1
            {filtro_giorno}
            ORDER BY c.nome, cg.ordine, cg.giorno, cg.orario
        """

        return pd.read_sql(
            query,
            conn,
            params=tuple(params_extra)
        )

    if is_istruttore():

        istruttore_id = st.session_state.utente_id

        query = f"""
            SELECT DISTINCT
                c.*,
                cg.giorno AS giorno_lezione,
                cg.orario AS orario_lezione
            FROM corsi c
            JOIN corso_giorni cg
                ON cg.corso_id = c.id
            JOIN assegnazioni_istruttori ai
                ON ai.corso_id = c.id
            WHERE ai.istruttore_id = ?
            AND ai.attiva = 1
            AND c.attivo = 1
            {filtro_giorno}
            AND (
                ai.data_specifica IS NULL
                OR ai.data_specifica = ?
                OR ? IS NULL
            )
            ORDER BY c.nome, cg.ordine, cg.giorno, cg.orario
        """

        data_str = (
            str(data_evento)
            if data_evento is not None
            else None
        )

        params = (
            [istruttore_id]
            + params_extra
            + [data_str, data_str]
        )

        return pd.read_sql(
            query,
            conn,
            params=tuple(params)
        )

    return pd.DataFrame()


# ============================================================
# FUNZIONI BAMBINI
# ============================================================
def aggiungi_bambino(
    nome,
    cognome,
    data_nascita,
    corso_id,
    email_genitore,
    telefono_genitore,
    note
):

    c.execute(
        """
        INSERT INTO bambini(
            nome,
            cognome,
            corso_id,
            email_genitore,
            telefono_genitore,
            note,
            attivo
        )
        VALUES(?,?,?,?,?,?,?,1)
        """,
        (
            nome,
            cognome,
            data_nascita,
            corso_id,
            email_genitore,
            telefono_genitore,
            note
        )
    )

    bambino_id = c.lastrowid

    conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )

    except Exception as e:

        print(
            f"Errore backup GitHub: {e}"
        )

    return bambino_id

def aggiorna_bambino(
    bambino_id,
    nome,
    cognome,
    data_nascita,
    corso_id,
    email_genitore,
    telefono_genitore,
    note,
    attivo
):

    c.execute(
        """
        UPDATE bambini
        SET
            nome = ?,
            cognome = ?,
            corso_id = ?,
            email_genitore = ?,
            telefono_genitore = ?,
            note = ?,
            attivo = ?
        WHERE id = ?
        """,
        (
            nome,
            cognome,
            corso_id,
            email_genitore,
            telefono_genitore,
            note,
            int(attivo),
            bambino_id
        )
    )

    conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )

    except Exception as e:

        print(
            f"Errore backup GitHub: {e}"
        )
            
def invia_credenziali_genitore_email(
    email_destinatario,
    nome_bambino,
    password
):
    
    mittente = st.secrets["GMAIL_ADDRESS"]
    password_mittente = st.secrets["GMAIL_PASSWORD"]
    
    app_url = st.secrets.get(
        "APP_URL",
        "link dell'app Streamlit"
    )
    
    testo = f"""
        Ciao,
    
        è stato creato il tuo accesso all'Area Genitore della Power Team Messina.
    
        Credenziali di accesso:
    
        Email: {email_destinatario}
        Password: {password}
    
        Link app:
        {app_url}
    
        Ti consigliamo di conservare queste credenziali con attenzione.
    
        Power Team Messina
    """
    
    msg = MIMEText(
        testo,
        "plain",
        "utf-8"
    )
    
    msg["Subject"] = "Credenziali accesso Corsi Nuoto"
    msg["From"] = mittente
    msg["To"] = email_destinatario
    
    server = smtplib.SMTP(
        "smtp.gmail.com",
        587
    )
    
    server.starttls()
    
    server.login(
        mittente,
        password_mittente
    )
    
    server.send_message(msg)
    
    server.quit()    

def elimina_bambino(bambino_id):

    c.execute(
        """
        DELETE FROM presenze
        WHERE bambino_id = ?
        """,
        (bambino_id,)
    )

    c.execute(
        """
        DELETE FROM bambini
        WHERE id = ?
        """,
        (bambino_id,)
    )

    conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )
    
    except Exception as e:
    
        print(
            f"Errore backup GitHub: {e}"
        )


# ============================================================
# FUNZIONI PRESENZE
# ============================================================
def ottieni_nome_stagione():

    oggi = datetime.now()

    if oggi.month >= 9:
        return f"{oggi.year}-{oggi.year+1}"

    return f"{oggi.year-1}-{oggi.year}"

def carica_file_drive(
    uploaded_file,
    folder_id,
    nome_file
):

    if uploaded_file is None:

        st.warning(
            f"{nome_file} è None"
        )

        return

    st.success(
        f"Caricamento {nome_file}"
    )

    st.write(
        "Nome originale:",
        uploaded_file.name
    )

    st.write(
        "Tipo:",
        uploaded_file.type
    )

    st.write(
        "Dimensione:",
        uploaded_file.size
    )

    service = get_drive_service()

    media = MediaIoBaseUpload(
        io.BytesIO(
            uploaded_file.getvalue()
        ),
        mimetype=uploaded_file.type,
        resumable=False
    )

    metadata = {
        "name": nome_file,
        "parents": [folder_id]
    }

    try:

        file = service.files().create(
            body=metadata,
            media_body=media,
            fields="id,name"
        ).execute()

        st.success(
            f"Caricato: {file['name']}"
        )

    except Exception as e:

        st.error(
            f"Errore upload {nome_file}"
        )

        st.code(
            str(e)
        )

        raise

def stagione_corrente():

    oggi = datetime.now()

    if oggi.month >= 9:

        return (
            f"{oggi.year}-"
            f"{oggi.year + 1}"
        )

    return (
        f"{oggi.year - 1}-"
        f"{oggi.year}"
    )
    
def get_presenze_corso_data(corso_id, data_evento):

    return pd.read_sql(
        """
        SELECT *
        FROM presenze
        WHERE corso_id = ?
        AND data = ?
        """,
        conn,
        params=(
            corso_id,
            str(data_evento)
        )
    )

def get_presenze_bambino(
    bambino_id
):

    return pd.read_sql(
        """
        SELECT *
        FROM presenze
        WHERE bambino_id = ?
        ORDER BY data
        """,
        conn,
        params=(bambino_id,)
    )

def valore_pdf(
    valore
):

    if valore is None:
        return ""

    try:

        if pd.isna(
            valore
        ):
            return ""

    except Exception:
        pass

    return str(
        valore
    )


def data_pdf(
    valore
):

    if valore is None:
        return ""

    try:

        return pd.to_datetime(
            valore,
            errors="coerce"
        ).strftime(
            "%d/%m/%Y"
        )

    except Exception:

        return str(
            valore
        )

def genera_pdf_compilato(
    bambino,
    scheda
):

    b = bambino.iloc[0]
    s = scheda.iloc[0]

    packet = io.BytesIO()

    c = canvas.Canvas(
        packet,
        pagesize=A4
    )

    c.setFont(
        "Helvetica",
        30
    )
    
    # =====================================================
    # DATI BAMBINO
    # =====================================================
    
    scrivi_centrato(
        c,
        f"{valore_pdf(b['nome'])} {valore_pdf(b['cognome'])}",
        165,
        693
    )
    
    scrivi_centrato(
        c,
        valore_pdf(
            s.get("luogo_nascita_bambino", "")
        ),
        400,
        693
    )
    
    scrivi_centrato(
        c,
        data_pdf(
            s.get("data_nascita_bambino", "")
        ),
        90,
        673
    )
    
    scrivi_centrato(
        c,
        valore_pdf(
            s.get("cf_bambino", "")
        ),
        420,
        673
    )
    
    scrivi_centrato(
        c,
        valore_pdf(
            s.get("comune_bambino", "")
        ),
        120,
        654
    )
    
    scrivi_centrato(
        c,
        valore_pdf(
            s.get("indirizzo_bambino", "")
        ),
        420,
        654
    )
    
    scrivi_centrato(
        c,
        valore_pdf(
            b.get("email_genitore", "")
        ),
        180,
        634
    )
    
    scrivi_centrato(
        c,
        valore_pdf(
            b.get("telefono_genitore", "")
        ),
        450,
        634
    )
    
    # =====================================================
    # DATA MODULO
    # =====================================================
    
    c.setFillColorRGB(
        1,
        1,
        1
    )
    
    c.rect(
        72,
        356,
        95,
        12,
        stroke=0,
        fill=1
    )
    
    c.setFillColorRGB(
        0,
        0,
        0
    )
    
    scrivi_centrato(
        c,
        datetime.now().strftime("%d/%m/%Y"),
        100,
        460
    )
    
    # =====================================================
    # DATI GENITORE
    # =====================================================
    
    scrivi_centrato(
        c,
        f"{valore_pdf(s.get('nome_genitore',''))} {valore_pdf(s.get('cognome_genitore',''))}",
        120,
        400
    )
    
    scrivi_centrato(
        c,
        valore_pdf(
            s.get("luogo_nascita_genitore", "")
        ),
        400,
        400
    )
    
    scrivi_centrato(
        c,
        data_pdf(
            s.get("data_nascita_genitore", "")
        ),
        90,
        380
    )
    
    scrivi_centrato(
        c,
        valore_pdf(
            s.get("cf_genitore", "")
        ),
        420,
        380
    )
    
    scrivi_centrato(
        c,
        f"{valore_pdf(s.get('indirizzo_genitore',''))} - {valore_pdf(s.get('comune_genitore',''))}",
        160,
        360
    )
    
    scrivi_centrato(
        c,
        valore_pdf(
            b.get("telefono_genitore", "")
        ),
        410,
        360
    )
    
    scrivi_centrato(
        c,
        valore_pdf(
            b.get("email_genitore", "")
        ),
        160,
        340
    )

    c.save()

    packet.seek(0)
    
    overlay_pdf = PdfReader(packet)
    
    base_pdf = PdfReader(
        "templates/modulo_base.pdf"
    )
    
    writer = PdfWriter()
    
    page = base_pdf.pages[0]
    
    page.merge_page(
        overlay_pdf.pages[0]
    )
    
    writer.add_page(page)
    
    output = io.BytesIO()
    
    writer.write(output)
    
    output.seek(0)
    
    return output

def salva_presenza(bambino_id, corso_id, data_evento, presenza, note):

    c.execute(
        """
        INSERT INTO presenze(
            bambino_id,
            corso_id,
            data,
            presenza,
            note,
            inserito_da
        )
        VALUES(?,?,?,?,?,?)

        ON CONFLICT(
            bambino_id,
            corso_id,
            data
        )

        DO UPDATE SET
            presenza = excluded.presenza,
            note = excluded.note,
            inserito_da = excluded.inserito_da
        """,
        (
            bambino_id,
            corso_id,
            str(data_evento),
            int(presenza),
            note,
            st.session_state.utente_id
        )
    )

    conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )
    
    except Exception as e:
    
        print(
            f"Errore backup GitHub: {e}"
        )

def aggiungi_chiusura(
    data,
    descrizione,
    corso_id=None
):

    c.execute(
        """
        INSERT INTO chiusure(
            data,
            descrizione,
            corso_id,
            attiva
        )
        VALUES(?,?,?,1)
        """,
        (
            data,
            descrizione,
            corso_id
        )
    )

    conn.commit()

    try:
    
        upload_backup_github(
            mostra_messaggio=False
        )
    
    except:
        pass

def get_chiusure():

    return pd.read_sql(
        """
        SELECT *
        FROM chiusure
        WHERE attiva = 1
        ORDER BY data
        """,
        conn
    )

def elimina_chiusura(chiusura_id):

    c.execute(
        """
        DELETE FROM chiusure
        WHERE id = ?
        """,
        (chiusura_id,)
    )

    conn.commit()

    try:

        upload_backup_github(
            mostra_messaggio=False
        )

    except:
        pass

def storico_presenze():

    return pd.read_sql(
        """
        SELECT
            p.data,
            c.nome AS corso,
            c.giorno,
            c.orario,
            b.cognome,
            b.nome,
            p.presenza,
            p.note,
            u.nome AS inserito_da
        FROM presenze p
        JOIN bambini b
            ON b.id = p.bambino_id
        JOIN corsi c
            ON c.id = p.corso_id
        LEFT JOIN utenti u
            ON u.id = p.inserito_da
        ORDER BY p.data DESC, c.nome, b.cognome, b.nome
        """,
        conn
    )

def get_stagioni():

    df = pd.read_sql(
        """
        SELECT DISTINCT stagione
        FROM corsi
        WHERE stagione IS NOT NULL
        AND stagione != ''
        ORDER BY stagione DESC
        """,
        conn
    )

    return df["stagione"].tolist()

def get_genitori():

    return pd.read_sql(
        """
        SELECT
            u.id,
            u.nome,
            u.username,
            b.telefono_genitore,
            u.password_visibile,
            u.attivo,
            b.nome AS nome_bambino,
            b.cognome AS cognome_bambino
        FROM utenti u
        LEFT JOIN genitori_bambini gb
            ON gb.utente_id = u.id
        LEFT JOIN bambini b
            ON b.id = gb.bambino_id
        WHERE u.ruolo = 'genitore'
        ORDER BY u.nome
        """,
        conn
    )

def get_bambino_associato_genitore(
    genitore_id
):

    return pd.read_sql(
        """
        SELECT
            b.nome,
            b.cognome
        FROM bambini b
        JOIN genitori_bambini gb
            ON gb.bambino_id = b.id
        WHERE gb.utente_id = ?
        """,
        conn,
        params=(genitore_id,)
    )
    
def backup_giornaliero():

    oggi = datetime.today().strftime(
        "%Y-%m-%d"
    )

    df = pd.read_sql(
        """
        SELECT *
        FROM sistema
        WHERE chiave='ultimo_backup'
        """,
        conn
    )

    if (
        df.empty
        or
        df.iloc[0]["valore"] != oggi
    ):

        crea_backup()

        c.execute(
            """
            INSERT OR REPLACE INTO sistema
            VALUES(
                'ultimo_backup',
                ?
            )
            """,
            (oggi,)
        )

        conn.commit()

        try:
    
            upload_backup_github(
                mostra_messaggio=False
            )
        
        except Exception as e:
        
            print(
                f"Errore backup GitHub: {e}"
            )

def get_bambino_genitore(
    utente_id
):

    return pd.read_sql(
        """
        SELECT b.*
        FROM bambini b
        JOIN genitori_bambini gb
            ON gb.bambino_id = b.id
        WHERE gb.utente_id = ?
        """,
        conn,
        params=(utente_id,)
    )
    
# ============================================================
# INTERFACCIA
# ============================================================

ripristino_iniziale_da_github()

BACKUP_ABILITATO = True

migra_giorni_corsi_vecchi()

crea_manager_default()

st.title("🏊 Statino Presenze Corsi di Nuoto")

stagioni = get_stagioni()

if len(stagioni) > 0:

    stagione_selezionata = st.selectbox(
        "Stagione",
        stagioni,
        key="stagione_corrente"
    )

else:

    stagione_selezionata = None

login()

if not is_loggato():

    st.info(
        "Effettua il login per accedere al gestionale."
    )

    st.stop()


st.markdown(
    f"""
    **Utente:** {st.session_state.nome_utente}  
    **Ruolo:** {st.session_state.ruolo}
    """
)


# ============================================================
# TAB
# ============================================================

if is_manager():

    tabs = st.tabs(
        [
            "📋 Presenze",
            "👶 Bambini",
            "🏊 Corsi",
            "📅 Stagioni",
            "👨‍🏫 Istruttori",
            "👔 Manager",
            "👨‍👩‍👧 Genitori",
            "🔗 Assegnazioni",
            "📄 Documentazione",
            "📊 Riepilogo presenze",
            "🚫 Chiusure",
            "🗂️ Storico",
            "💾 Backup"
        ]
    )

    tab_presenze = tabs[0]
    tab_bambini = tabs[1]
    tab_corsi = tabs[2]
    tab_stagioni = tabs[3]
    tab_istruttori = tabs[4]
    tab_manager = tabs[5]
    tab_genitori = tabs[6]
    tab_assegnazioni = tabs[7]
    tab_documentazione = tabs[8]
    tab_riepilogo = tabs[9]
    tab_chiusure = tabs[10]
    tab_storico = tabs[11]
    tab_backup = tabs[12]

elif is_genitore():

    tabs = st.tabs(
        [
            "🏠 Area personale"
        ]
    )

    tab_area_personale = tabs[0]

else:

    tabs = st.tabs(
        [
            "📋 Presenze",
            "👶 Bambini"
        ]
    )

    tab_presenze = tabs[0]
    tab_bambini = tabs[1]

# ============================================================
# TAB PRESENZE
# ============================================================

if not is_genitore():
    with tab_presenze:

        st.header("📋 Registro presenze")
    
        data_evento = st.date_input(
            "Data",
            value=date.today(),
            key="data_presenze"
        )
    
        corsi_visibili = get_corsi_visibili_per_utente(
            data_evento
        )
    
        if corsi_visibili.empty:
    
            st.warning(
                "Non hai corsi disponibili per questa data."
            )
    
        else:
    
            corsi_visibili = get_corsi_con_giorni(
                attivi_solo=True
            ) if is_manager() else get_corsi_visibili_per_utente()
            
            opzioni_corsi = {
                f"{row['nome']} | {row['giorni_orari']}": int(row["id"])
                for _, row in get_corsi_con_giorni(
                    attivi_solo=True
                ).iterrows()
            }
    
            corso_label = st.selectbox(
                "Corso",
                list(opzioni_corsi.keys()),
                key="corso_presenze"
            )
    
            corso_id = opzioni_corsi[corso_label]
    
            if is_istruttore():
    
                abilitato = istruttore_abilitato_corso_data(
                    st.session_state.utente_id,
                    corso_id,
                    data_evento
                )
    
                if not abilitato:
    
                    st.error(
                        "Non sei abilitato a compilare questo corso in questa data."
                    )
    
                    st.stop()
    
            bambini = get_bambini_corso(
                corso_id,
                attivi_solo=True
            )
    
            if bambini.empty:
    
                st.info(
                    "Nessun bambino inserito in questo corso."
                )
    
            else:
    
                presenze_salvate = get_presenze_corso_data(
                    corso_id,
                    data_evento
                )
    
                dati_salvati = {}
    
                for _, r in presenze_salvate.iterrows():
    
                    dati_salvati[int(r["bambino_id"])] = {
                        "presenza": bool(r["presenza"]),
                        "note": r["note"] if pd.notna(r["note"]) else ""
                    }
    
                st.markdown("---")
    
                presenti = 0
                registro = {}
    
                for _, b in bambini.iterrows():
    
                    bambino_id = int(b["id"])
    
                    default_presenza = dati_salvati.get(
                        bambino_id,
                        {}
                    ).get(
                        "presenza",
                        False
                    )
    
                    default_note = dati_salvati.get(
                        bambino_id,
                        {}
                    ).get(
                        "note",
                        ""
                    )
    
                    st.subheader(
                        f"{b['cognome']} {b['nome']}"
                    )
    
                    col1, col2 = st.columns(
                        [1, 3]
                    )
    
                    presenza = col1.toggle(
                        "Presente",
                        value=default_presenza,
                        key=f"pres_{corso_id}_{data_evento}_{bambino_id}"
                    )
    
                    note = col2.text_input(
                        "Note",
                        value=default_note,
                        key=f"note_{corso_id}_{data_evento}_{bambino_id}"
                    )
    
                    if presenza:
                        presenti += 1
    
                    registro[bambino_id] = {
                        "presenza": presenza,
                        "note": note
                    }
    
                    st.markdown("---")
    
                assenti = len(bambini) - presenti
    
                c1, c2, c3 = st.columns(3)
    
                c1.metric(
                    "Iscritti",
                    len(bambini)
                )
    
                c2.metric(
                    "Presenti",
                    presenti
                )
    
                c3.metric(
                    "Assenti",
                    assenti
                )
    
                if st.button(
                    "💾 Salva presenze",
                    key="salva_presenze"
                ):
    
                    for bambino_id, dati in registro.items():
    
                        salva_presenza(
                            bambino_id,
                            corso_id,
                            data_evento,
                            dati["presenza"],
                            dati["note"]
                        )
    
                    st.success(
                        "Presenze salvate correttamente."
                    )
    
                    st.rerun()
    
                st.markdown("---")
    
                conferma_elimina_presenze = st.checkbox(
                    f"Confermo eliminazione presenze del {data_evento}",
                    key=f"conferma_elimina_presenze_{corso_id}_{data_evento}"
                )
    
                presenze_presenti = get_presenze_corso_data(
                    corso_id,
                    data_evento
                )
                
                st.info(
                    f"Attualmente sono registrate {len(presenze_presenti)} presenze per questo corso in questa data."
                )
                
                if st.button(
                    "🗑️ Elimina presenze del giorno",
                    key=f"elimina_presenze_{corso_id}_{data_evento}"
                ):
                
                    if not conferma_elimina_presenze:
                
                        st.error(
                            "Devi confermare l'eliminazione."
                        )
                
                    else:
                
                        numero = elimina_presenze_giornata(
                            corso_id,
                            data_evento
                        )
                        
                        st.success(
                            f"{numero} presenze eliminate."
                        )
                
                        st.rerun()


# ============================================================
# TAB BAMBINI
# ============================================================

if not is_genitore():
    with tab_bambini:

        if is_manager():
    
            st.header("👶 Gestione bambini")
    
            corsi = get_corsi(
                attivi_solo=True,
                stagione=stagione_selezionata
            )
            
            opzioni_corsi = {
                f"{row['nome']}":
                    int(row["id"])
                for _, row in corsi.iterrows()
            }
        
            st.subheader("➕ Aggiungi bambino")
        
            with st.form(
                "form_aggiungi_bambino",
                clear_on_submit=True
            ):
        
                nome = st.text_input(
                    "Nome"
                )
        
                cognome = st.text_input(
                    "Cognome"
                )
                
                email_genitore = st.text_input(
                    "Email genitore (facoltativa)"
                )
    
                telefono_genitore = st.text_input(
                    "Telefono genitore (facoltativo)"
                )
    
                corsi = get_corsi(
                    attivi_solo=True
                )
                
                if len(opzioni_corsi) == 0:
                
                    st.error(
                        "Non esistono corsi attivi. Crea prima almeno un corso."
                    )
                
                else:
                
                    corso_label = st.selectbox(
                        "Corso principale",
                        list(opzioni_corsi.keys())
                    )
                
                    corso_id = opzioni_corsi[
                        corso_label
                    ]
        
                note = st.text_area(
                    "Note"
                )
        
                invia = st.form_submit_button(
                    "➕ Aggiungi"
                )
        
                if invia:
        
                    if nome.strip() == "" or cognome.strip() == "":
        
                        st.error(
                            "Nome e cognome sono obbligatori."
                        )
        
                    else:
        
                        try:
    
                            bambino_id = aggiungi_bambino(
                                nome,
                                cognome,
                                corso_id,
                                email_genitore,
                                telefono_genitore,
                                note
                            )
                            
                            if email_genitore.strip() != "":
                            
                                genitore_esistente = pd.read_sql(
                                    """
                                    SELECT *
                                    FROM utenti
                                    WHERE lower(username)=?
                                    """,
                                    conn,
                                    params=(email_genitore.strip().lower(),)
                                )
                            
                                if genitore_esistente.empty:
                            
                                    
                                    genitore_id, password_generata, gia_esistente = (
                                        crea_account_genitore(
                                            dati["email_genitore"],
                                            f"{dati['nome']} {dati['cognome']}"
                                        )
                                    )
                            
                                    c.execute(
                                        """
                                        INSERT OR IGNORE INTO genitori_bambini(
                                            utente_id,
                                            bambino_id
                                        )
                                        VALUES(?,?)
                                        """,
                                        (
                                            genitore_id,
                                            bambino_id
                                        )
                                    )
                            
                                    conn.commit()
                            
                                    try:
                            
                                        invia_credenziali_genitore_email(
                                            email_genitore,
                                            f"{nome} {cognome}",
                                            password_generata
                                        )
                            
                                    except Exception as e:
                            
                                        st.warning(
                                            f"Account creato ma email non inviata: {e}"
                                        )
                        
                            st.success(
                                f"Bambino salvato. ID={bambino_id}"
                            )
                        
                        except Exception as e:
                        
                            st.error(str(e))
                            
                        st.success(
                            "Bambino aggiunto correttamente."
                        )
        
                        st.rerun()
    
        st.markdown("---")
    
        st.subheader("📋 Elenco bambini")
    
        bambini = get_bambini(
            attivi_solo=False if is_manager() else True
        )
    
        if bambini.empty:
    
            st.info(
                "Nessun bambino presente."
            )
    
        else:
    
            bambini_visual = bambini.copy()
    
            st.dataframe(
                bambini_visual[
                    [
                        "id",
                        "cognome",
                        "nome",
                        "email_genitore",
                        "telefono_genitore",
                        "note",
                        "attivo"
                    ]
                ],
                use_container_width=True,
                hide_index=True
            )
    
            if is_manager():
    
                st.markdown("---")
    
                st.subheader("✏️ Modifica o elimina bambino")
    
                opzioni_bambini = {
                    f"{row['cognome']} {row['nome']}": int(row["id"])
                    for _, row in bambini.iterrows()
                }
    
                bambino_label = st.selectbox(
                    "Bambino",
                    list(opzioni_bambini.keys()),
                    key="modifica_bambino"
                )
    
                bambino_id = opzioni_bambini[bambino_label]
    
                dati = bambini[
                    bambini["id"] == bambino_id
                ].iloc[0]
    
                nuovo_nome = st.text_input(
                    "Nome",
                    value=dati["nome"],
                    key="nuovo_nome_bambino"
                )
    
                nuovo_cognome = st.text_input(
                    "Cognome",
                    value=dati["cognome"],
                    key="nuovo_cognome_bambino"
                )
                    
                nuova_email_genitore = st.text_input(
                    "Email genitore",
                    value=dati["email_genitore"]
                    if pd.notna(dati["email_genitore"])
                    else ""
                )
    
                nuovo_telefono_genitore = st.text_input(
                    "Telefono genitore",
                    value=dati["telefono_genitore"]
                    if pd.notna(dati["telefono_genitore"])
                    else ""
                )
    
                nuove_note = st.text_area(
                    "Note",
                    value=dati["note"] if pd.notna(dati["note"]) else "",
                    key="nuove_note_bambino"
                )
    
                nuovo_attivo = st.checkbox(
                    "Attivo",
                    value=bool(dati["attivo"]),
                    key="attivo_bambino"
                )
    
                corsi = get_corsi(
                    attivi_solo=False
                )
                    
                opzioni_corsi = {
                    row["nome"]: int(row["id"])
                    for _, row in corsi.iterrows()
                }
    
                nuovo_corso = st.selectbox(
                    "Corso principale",
                    list(opzioni_corsi.keys())
                )
    
                if st.button(
                    "💾 Aggiorna bambino"
                ):
    
                    aggiorna_bambino(
                        bambino_id,
                        nuovo_nome.strip(),
                        nuovo_cognome.strip(),
                        opzioni_corsi[nuovo_corso],
                        nuova_email_genitore.strip(),
                        nuovo_telefono_genitore.strip(),
                        nuove_note.strip(),
                        nuovo_attivo
                    )
    
                    st.success(
                        "Bambino aggiornato."
                    )
    
                    st.rerun()
                    
            st.markdown("---")
            st.subheader("👨‍👩‍👧 Account Genitore")
    
            genitore = pd.read_sql(
                """
                SELECT u.*
                FROM utenti u
                JOIN genitori_bambini gb
                ON gb.utente_id = u.id
                WHERE gb.bambino_id = ?
                AND u.ruolo = 'genitore'
                """,
                conn,
                params=(bambino_id,)
            )
                    
            if not genitore.empty:
                    
                genitore_row = genitore.iloc[0]
                    
                st.text_input(
                    "Email account genitore",
                    value=genitore_row["username"],
                    disabled=True,
                    key=f"email_account_genitore_{bambino_id}"
                )
    
                st.text_input(
                    "Password corrente account",
                    value=genitore_row["password_visibile"]
                    if pd.notna(
                        genitore_row["password_visibile"]
                    )
                    else "",
                    disabled=True,
                    key=f"password_account_genitore_{bambino_id}"
                )
    
                if st.button(
                    "📧 Reinvia credenziali genitore"
                ):
                        
                    invia_credenziali_genitore_email(
                        genitore_row["username"],
                        f"{dati['nome']} {dati['cognome']}",
                        genitore_row["password_visibile"]
                    )
                        
                    st.success(
                        "Credenziali inviate."
                    )
    
                if st.button(
                    "🔄 Genera nuova password"
                ):
                        
                    nuova_password = genera_password_casuale()
                        
                    aggiorna_password_utente(
                        int(genitore_row["id"]),
                        nuova_password
                    )
                        
                    invia_credenziali_genitore_email(
                        genitore_row["username"],
                        f"{dati['nome']} {dati['cognome']}",
                        nuova_password
                    )
                        
                    st.success(
                        "Nuova password generata e inviata."
                    )
                        
                    st.rerun()
                    
            else:
                    
                st.warning(
                    "Nessun account genitore associato."
                )
    
                if (
                    pd.notna(dati["email_genitore"])
                    and
                    dati["email_genitore"].strip() != ""
                ):
                        
                    if st.button(
                         "➕ Crea account genitore"
                    ):
                        
                        genitore_id, password_generata, gia_esistente = (
                            crea_account_genitore(
                                dati["email_genitore"],
                                f"{dati['nome']} {dati['cognome']}"
                            )
                        )
    
                        if gia_esistente:
                            st.info("Account genitore già presente.")
                        else:
                            st.success(
                                f"Account creato. Password: {password_generata}"
                            )
                        
                        c.execute(
                            """
                            INSERT INTO genitori_bambini(
                                utente_id,
                                bambino_id
                            )
                            VALUES(?,?)
                            """,
                            (
                                genitore_id,
                                bambino_id
                            )
                        )
                        
                        conn.commit()
                    
                        st.success(
                            f"Account creato. Password: {password_generata}"
                        )
                    
                        st.rerun()
                    
                conferma_elimina = st.checkbox(
                    "Confermo eliminazione definitiva bambino"
                )
    
                if st.button(
                    "🗑️ Elimina bambino"
                ):
    
                    if not conferma_elimina:
    
                        st.error(
                            "Devi confermare l'eliminazione."
                        )
    
                    else:
    
                        elimina_bambino(
                            bambino_id
                        )
    
                        st.success(
                            "Bambino eliminato."
                        )
    
                        st.rerun()

######### STAGIONI #########
if is_manager():
    with tab_stagioni:

        st.header("📅 Gestione stagioni")
    
        nuova_stagione = st.text_input(
            "Nuova stagione",
            placeholder="es. 2027/2028"
        )
    
        if st.button("➕ Aggiungi stagione"):
    
            if nuova_stagione.strip() == "":
    
                st.error(
                    "Inserisci il nome della stagione."
                )
    
            else:
    
                try:
    
                    aggiungi_stagione(
                        nuova_stagione.strip()
                    )
    
                    st.success(
                        "Stagione aggiunta."
                    )
    
                    st.rerun()
    
                except:
    
                    st.error(
                        "Stagione già esistente."
                    )
    
        st.markdown("---")
    
        st.dataframe(
            get_stagioni(),
            use_container_width=True,
            hide_index=True
        )
    
        st.markdown("---")
    
        st.subheader("📊 Riepilogo stagioni")
        
        riepilogo_stagioni = get_riepilogo_stagioni()
        
        if riepilogo_stagioni.empty:
        
            st.info(
                "Nessuna stagione presente."
            )
        
        else:
        
            st.dataframe(
                riepilogo_stagioni.rename(
                    columns={
                        "stagione": "Stagione",
                        "numero_corsi": "Numero corsi",
                        "numero_istruttori": "Numero istruttori",
                        "numero_bambini": "Numero bambini"
                    }
                ),
                use_container_width=True,
                hide_index=True
            )

# ============================================================
# TAB CORSI
# ============================================================

if is_manager():

    with tab_corsi:

        st.header("🏊 Gestione corsi")

        giorni_settimana = [
            "Lunedì",
            "Martedì",
            "Mercoledì",
            "Giovedì",
            "Venerdì",
            "Sabato",
            "Domenica"
        ]

        # =====================================================
        # NUOVO CORSO
        # =====================================================

        st.subheader("➕ Nuovo corso")

        nome = st.text_input(
            "Nome corso",
            placeholder="es. Corso Bambini 1"
        )

        livello = st.text_input(
            "Livello",
            placeholder="es. Principianti"
        )

        stagione = st.selectbox(
            "Stagione",
            list(dict.fromkeys(
                ["2026/2027"] + get_stagioni()
            ))
        )

        giorni_selezionati = st.multiselect(
            "Giorni del corso",
            [
                "Lunedì",
                "Martedì",
                "Mercoledì",
                "Giovedì",
                "Venerdì",
                "Sabato",
                "Domenica"
            ],
            max_selections=5
        )
            
        giorni_orari = []
            
        for giorno in giorni_selezionati:
            
            orario = st.text_input(
                f"Orario {giorno}",
                placeholder="es. 16:00-17:00",
                key=f"nuovo_orario_{giorno}"
            )
            
            giorni_orari.append(
                (
                    giorno,
                    orario
                )
            )

        crea = st.button(
            "➕ Crea corso"
        )

        if crea:

            if nome.strip() == "":

                st.error(
                    "Inserisci il nome del corso."
                )

            elif len(giorni_orari) == 0:

                st.error(
                    "Seleziona almeno un giorno."
                )
                
            elif any(
                orario.strip() == ""
                for _, orario in giorni_orari
            ):
                
                st.error(
                    "Inserisci tutti gli orari."
                )
                
            elif len(
                set(giorno for giorno, _ in giorni_orari)
            ) != len(giorni_orari):
                
                st.error(
                    "Non puoi inserire due volte lo stesso giorno nello stesso corso."
                )
                
            else:
                
                corso_id = aggiungi_corso(
                    nome.strip(),
                    livello.strip(),
                    stagione.strip()
                )
                
                salva_giorni_corso(
                    corso_id,
                    giorni_orari
                )
                
                st.success(
                    "Corso creato correttamente."
                )
                
                st.rerun()

        st.markdown("---")

        # =====================================================
        # LISTA CORSI
        # =====================================================

        st.subheader("📋 Corsi presenti")

        corsi = get_corsi_con_giorni(
            attivi_solo=True
        )

        if corsi.empty:

            st.info(
                "Nessun corso presente."
            )

        else:

            st.dataframe(
                corsi[
                    [
                        "id",
                        "nome",
                        "livello",
                        "stagione",
                        "giorni_orari",
                        "attivo"
                    ]
                ],
                use_container_width=True,
                hide_index=True
            )

            st.markdown("---")

            # =====================================================
            # MODIFICA CORSO
            # =====================================================

            st.subheader("✏️ Modifica corso")

            opzioni_corsi = {
                f"{row['nome']} | {row['giorni_orari']}": int(row["id"])
                for _, row in corsi.iterrows()
            }

            corso_label = st.selectbox(
                "Seleziona corso",
                list(opzioni_corsi.keys()),
                key="modifica_corso"
            )

            corso_id = opzioni_corsi[corso_label]

            dati_corso = corsi[
                corsi["id"] == corso_id
            ].iloc[0]

            nuovo_nome = st.text_input(
                "Nome corso",
                value=dati_corso["nome"],
                key="mod_nome_corso"
            )

            nuovo_livello = st.text_input(
                "Livello",
                value=dati_corso["livello"] if pd.notna(dati_corso["livello"]) else "",
                key="mod_livello_corso"
            )

            nuova_stagione = st.text_input(
                "Stagione",
                value=dati_corso["stagione"] if pd.notna(dati_corso["stagione"]) else "",
                key="mod_stagione_corso"
            )

            nuovo_attivo = st.checkbox(
                "Corso attivo",
                value=bool(dati_corso["attivo"]),
                key="mod_attivo_corso"
            )

            giorni_attuali = get_giorni_corso(
                corso_id
            )

            numero_giorni_attuale = max(
                1,
                min(
                    5,
                    len(giorni_attuali)
                )
            )

            nuovo_numero_giorni = st.number_input(
                "Numero giorni settimanali",
                min_value=1,
                max_value=5,
                value=numero_giorni_attuale,
                step=1,
                key="mod_numero_giorni"
            )

            nuovi_giorni_orari = []

            st.markdown("### Giorni e orari del corso")

            giorni_settimana = [
                "Lunedì",
                "Martedì",
                "Mercoledì",
                "Giovedì",
                "Venerdì",
                "Sabato",
                "Domenica"
            ]
            
            giorni_predefiniti = (
                giorni_attuali["giorno"].tolist()
                if not giorni_attuali.empty
                else []
            )
            
            giorni_selezionati = st.multiselect(
                "Giorni del corso",
                giorni_settimana,
                default=giorni_predefiniti,
                max_selections=5,
                key=f"giorni_corso_{corso_id}"
            )
            
            nuovi_giorni_orari = []
            
            for giorno in giorni_selezionati:
            
                riga = giorni_attuali[
                    giorni_attuali["giorno"] == giorno
                ]
            
                if not riga.empty:
            
                    orario_default = (
                        riga.iloc[0]["orario"]
                    )
            
                else:
            
                    orario_default = ""
            
                orario = st.text_input(
                    f"Orario {giorno}",
                    value=orario_default,
                    key=f"orario_corso_{corso_id}_{giorno}"
                )
            
                nuovi_giorni_orari.append(
                    (
                        giorno,
                        orario
                    )
                )

            if st.button(
                "💾 Aggiorna corso"
            ):

                if nuovo_nome.strip() == "":

                    st.error(
                        "Il nome del corso non può essere vuoto."
                    )
                
                elif len(nuovi_giorni_orari) == 0:
                
                    st.error(
                        "Seleziona almeno un giorno."
                    )
                
                elif any(
                    orario.strip() == ""
                    for _, orario in nuovi_giorni_orari
                ):
                
                    st.error(
                        "Inserisci tutti gli orari."
                    )
                
                elif len(
                    set(giorno for giorno, _ in nuovi_giorni_orari)
                ) != len(nuovi_giorni_orari):
                
                    st.error(
                        "Non puoi inserire due volte lo stesso giorno nello stesso corso."
                    )
                
                else:
                
                    aggiorna_corso(
                        corso_id,
                        nuovo_nome.strip(),
                        nuovo_livello.strip(),
                        nuova_stagione.strip(),
                        nuovo_attivo
                    )
                
                    salva_giorni_corso(
                        corso_id,
                        nuovi_giorni_orari
                    )
                
                    st.success(
                        "Corso aggiornato correttamente."
                    )
                
                    st.rerun()

            st.markdown("---")

            # =====================================================
            # ELIMINA CORSO
            # =====================================================

            st.subheader("🗑️ Elimina corso")

            conferma_elimina_corso = st.checkbox(
                "Confermo eliminazione definitiva corso"
            )

            if st.button(
                "🗑️ Elimina corso"
            ):

                if not conferma_elimina_corso:

                    st.error(
                        "Devi confermare l'eliminazione."
                    )

                else:

                    elimina_corso(
                        corso_id
                    )

                    st.success(
                        "Corso eliminato."
                    )

                    st.rerun()


# ============================================================
# TAB ISTRUTTORI
# ============================================================

if is_manager():

    with tab_istruttori:

        st.header("👨‍🏫 Gestione istruttori")

        st.subheader("➕ Nuovo istruttore")

        with st.form(
            "form_nuovo_istruttore",
            clear_on_submit=True
        ):
        
            email = st.text_input(
                "Email istruttore"
            )
        
            nome = st.text_input(
                "Nome istruttore"
            )
        
            crea = st.form_submit_button(
                "➕ Crea istruttore e invia credenziali"
            )
        
            if crea:
        
                if (
                    email.strip() == ""
                    or nome.strip() == ""
                ):
        
                    st.error(
                        "Compila email e nome istruttore."
                    )
        
                elif "@" not in email.strip():
        
                    st.error(
                        "Inserisci un indirizzo email valido."
                    )
        
                else:
        
                    try:
        
                        password_generata = aggiungi_istruttore(
                            email.strip(),
                            nome.strip()
                        )
        
                        try:
        
                            invia_credenziali_istruttore_email(
                                email.strip().lower(),
                                nome.strip(),
                                password_generata
                            )
        
                            st.success(
                                "Istruttore creato e credenziali inviate via email."
                            )
        
                        except Exception as e:
        
                            st.warning(
                                "Istruttore creato, ma invio email non riuscito."
                            )
        
                            st.info(
                                f"Password generata da comunicare manualmente: {password_generata}"
                            )
        
                            st.error(
                                f"Errore email: {e}"
                            )
        
                        st.rerun()
        
                    except sqlite3.IntegrityError:
        
                        st.error(
                            "Esiste già un istruttore con questa email."
                        )

        st.markdown("---")

        st.subheader("📋 Istruttori")

        istruttori = get_istruttori(
            attivi_solo=False
        )

        if istruttori.empty:

            st.info(
                "Nessun istruttore presente."
            )

        else:

            istruttori_visibili = pd.read_sql(
                """
                SELECT
                    id,
                    nome,
                    username AS email,
                    password_visibile,
                    attivo
                FROM utenti
                WHERE ruolo='istruttore'
                ORDER BY nome
                """,
                conn
            )
            
            st.dataframe(
                istruttori_visibili,
                use_container_width=True,
                hide_index=True
            )

            st.markdown("---")

            st.subheader("⚙️ Modifica istruttore")

            opzioni = {
                f"{row['nome']} ({row['username']})": int(row["id"])
                for _, row in istruttori.iterrows()
            }

            scelta = st.selectbox(
                "Istruttore",
                list(opzioni.keys()),
                key="istruttore_modifica"
            )

            istruttore_id = opzioni[scelta]

            dati_istruttore = istruttori[
                istruttori["id"] == istruttore_id
            ].iloc[0]

            dettagli = pd.read_sql(
                """
                SELECT *
                FROM utenti
                WHERE id = ?
                """,
                conn,
                params=(istruttore_id,)
            ).iloc[0]

            attivo = st.checkbox(
                "Account attivo",
                value=bool(dati_istruttore["attivo"]),
                key="stato_istruttore"
            )

            st.text_input(
                "Email",
                value=dettagli["username"],
                disabled=True
            )

            st.text_input(
                "Password corrente",
                value=dettagli["password_visibile"]
                    if pd.notna(
                        dettagli["password_visibile"]
                    )
                    else "",
                disabled=True
            )

            if st.button(
                "📧 Reinvia credenziali"
            ):
            
                invia_credenziali_istruttore_email(
                    dettagli["username"],
                    dettagli["nome"],
                    dettagli["password_visibile"]
                )
            
                st.success(
                    "Credenziali inviate."
                )

            if st.button(
                "💾 Aggiorna stato account"
            ):

                cambia_stato_utente(
                    istruttore_id,
                    attivo
                )

                st.success(
                    "Stato account aggiornato."
                )

                st.rerun()

            nuova_password = st.text_input(
                "Nuova password",
                type="password",
                key="nuova_password_istruttore"
            )

            if st.button(
                "🔐 Cambia password istruttore"
            ):

                if nuova_password.strip() == "":

                    st.error(
                        "Inserisci una nuova password."
                    )

                else:

                    aggiorna_password_utente(
                        istruttore_id,
                        nuova_password.strip()
                    )

                    st.success(
                        "Password aggiornata."
                    )

            st.markdown("---")

            st.write("DEBUG ELIMINAZIONE")
            
            st.subheader("🗑️ Elimina istruttore")
            
            conferma_elimina = st.checkbox(
                f"Confermo eliminazione di {dettagli['nome']}",
                key=f"elimina_{istruttore_id}"
            )
                
            if st.button(
                "🗑️ Elimina istruttore",
                key=f"btn_elimina_{istruttore_id}"
            ):
                
                if not conferma_elimina:
                
                    st.error(
                        "Devi confermare l'eliminazione."
                    )
            
                else:
                
                    elimina_istruttore(
                        istruttore_id
                    )
            
                    st.success(
                        "Istruttore eliminato."
                    )
            
                    st.rerun()


# ============================================================
# TAB ASSEGNAZIONI
# ============================================================

if is_manager():

    with tab_assegnazioni:

        st.header("🔗 Assegnazione istruttori ai corsi")

        istruttori = get_istruttori(
            attivi_solo=True
        )

        get_corsi(
            attivi_solo=False,
            stagione=stagione_selezionata
        )

        if istruttori.empty or corsi.empty:

            st.info(
                "Servono almeno un istruttore attivo e un corso attivo."
            )

        else:

            opzioni_istruttori = {
                f"{row['nome']} ({row['username']})": int(row["id"])
                for _, row in istruttori.iterrows()
            }

            corsi_assegnabili = get_corsi_con_giorni(
                attivi_solo=True
            )
            
            opzioni_corsi = {
                f"{row['nome']} | {row['giorni_orari']}": int(row["id"])
                for _, row in corsi_assegnabili.iterrows()
            }

            istruttore_label = st.selectbox(
                "Istruttore",
                list(opzioni_istruttori.keys()),
                key="assegna_istruttore"
            )

            corso_label = st.selectbox(
                "Corso",
                list(opzioni_corsi.keys()),
                key="assegna_corso"
            )

            tipo_assegnazione = st.radio(
                "Tipo assegnazione",
                [
                    "Intero corso",
                    "Solo una data specifica"
                ],
                horizontal=True
            )

            data_specifica = None

            if tipo_assegnazione == "Solo una data specifica":

                data_specifica = st.date_input(
                    "Data specifica",
                    value=date.today(),
                    key="data_specifica_assegnazione"
                )

            if st.button(
                "➕ Assegna istruttore"
            ):

                assegna_istruttore(
                    opzioni_istruttori[istruttore_label],
                    opzioni_corsi[corso_label],
                    str(data_specifica) if data_specifica else None
                )

                st.success(
                    "Assegnazione inserita."
                )

                st.rerun()

        st.markdown("---")

        st.subheader("📋 Assegnazioni presenti")

        assegnazioni = get_assegnazioni()

        if assegnazioni.empty:

            st.info(
                "Nessuna assegnazione presente."
            )

        else:

            st.dataframe(
                assegnazioni,
                use_container_width=True,
                hide_index=True
            )

            opzioni_ass = {
                f"{row['id']} | {row['istruttore']} | {row['corso']} | {row['data_specifica']}": int(row["id"])
                for _, row in assegnazioni.iterrows()
            }

            scelta_ass = st.selectbox(
                "Assegnazione da eliminare",
                list(opzioni_ass.keys()),
                key="elimina_assegnazione"
            )

            conferma = st.checkbox(
                "Confermo eliminazione assegnazione"
            )

            if st.button(
                "🗑️ Elimina assegnazione"
            ):

                if not conferma:

                    st.error(
                        "Devi confermare."
                    )

                else:

                    elimina_assegnazione(
                        opzioni_ass[scelta_ass]
                    )

                    st.success(
                        "Assegnazione eliminata."
                    )

                    st.rerun()


# ============================================================
# TAB STORICO
# ============================================================

if is_manager():
    with tab_storico:

        st.header("🗂️ Storico presenze")
    
        storico = storico_presenze()
    
        if is_istruttore():
    
            corsi_visibili = get_corsi_visibili_per_utente()
    
            ids_corsi = corsi_visibili["id"].tolist()
    
            storico = storico[
                storico["corso"].isin(
                    corsi_visibili["nome"].tolist()
                )
            ]
    
        if storico.empty:
    
            st.info(
                "Nessuna presenza registrata."
            )
    
        else:
    
            storico["presenza"] = storico["presenza"].map(
                {
                    1: "Presente",
                    0: "Assente"
                }
            )
    
            st.dataframe(
                storico,
                use_container_width=True,
                hide_index=True
            )
    
            csv = storico.to_csv(
                index=False
            ).encode(
                "utf-8"
            )
    
            st.download_button(
                "📥 Scarica storico CSV",
                csv,
                "storico_presenze_corsi.csv",
                "text/csv"
            )
            
        if is_manager():
    
            with tab_backup:
        
                st.header("💾 Backup GitHub")
        
                st.info(
                    "Il backup contiene istruttori, password, corsi, bambini, assegnazioni, presenze e stagioni."
                )
        
                if st.button(
                    "💾 Salva backup su GitHub"
                ):
        
                    upload_backup_github(
                        mostra_messaggio=True
                    )
        
                st.markdown("---")
        
                if os.path.exists(
                    "backup_completo.json"
                ):
        
                    with open(
                        "backup_completo.json",
                        "rb"
                    ) as f:
        
                        st.download_button(
                            "📥 Scarica backup JSON",
                            f,
                            file_name="backup_completo.json",
                            mime="application/json"
                        )
        
                st.markdown("---")
        
                st.subheader("📤 Carica backup manuale")
        
                uploaded_file = st.file_uploader(
                    "Carica file backup_completo.json",
                    type=["json"]
                )
        
                if uploaded_file is not None:
        
                    with open(
                        "backup_completo.json",
                        "wb"
                    ) as f:
        
                        f.write(
                            uploaded_file.getbuffer()
                        )
        
                    if st.button(
                        "♻️ Ripristina backup caricato"
                    ):
        
                        if ripristina_backup_locale():
        
                            st.success(
                                "Backup ripristinato correttamente."
                            )
        
                            upload_backup_github(
                                mostra_messaggio=True
                            )
        
                            st.rerun()
        
                        else:
        
                            st.error(
                                "Impossibile ripristinare il backup."
                            )
        
                st.markdown("---")
        
                st.subheader("☁️ Ripristino da GitHub")
        
                if st.button(
                    "📥 Scarica e ripristina backup da GitHub"
                ):
        
                    if scarica_backup_github():
        
                        ripristina_backup_locale()
        
                        st.success(
                            "Backup scaricato da GitHub e ripristinato."
                        )
        
                        st.rerun()
        
                    else:
        
                        st.error(
                            "Nessun backup trovato su GitHub oppure accesso non riuscito."
                        )

if is_manager():

    with tab_manager:

        st.header("👔 Gestione Manager")

        st.subheader("➕ Nuovo manager")

        with st.form(
            "form_nuovo_manager",
            clear_on_submit=True
        ):

            email = st.text_input(
                "Email manager"
            )

            nome = st.text_input(
                "Nome manager"
            )

            crea = st.form_submit_button(
                "➕ Crea manager e invia credenziali"
            )

            if crea:

                try:

                    password_generata = aggiungi_manager(
                        email,
                        nome
                    )

                    invia_credenziali_manager_email(
                        email,
                        nome,
                        password_generata
                    )

                    st.success(
                        "Manager creato correttamente."
                    )

                    st.rerun()

                except sqlite3.IntegrityError:

                    st.error(
                        "Email già presente."
                    )
        
        manager_df = pd.read_sql(
            """
            SELECT
                id,
                nome,
                username AS email,
                password_visibile,
                attivo
            FROM utenti
            WHERE ruolo='manager'
            ORDER BY nome
            """,
            conn
        )
        
        st.dataframe(
            manager_df,
            use_container_width=True,
            hide_index=True
        )

        st.markdown("---")

        st.subheader("⚙️ Modifica manager")

        opzioni_manager = {
            f"{row['nome']} ({row['email']})": int(row["id"])
            for _, row in manager_df.iterrows()
        }
        
        scelta_manager = st.selectbox(
            "Manager",
            list(opzioni_manager.keys()),
            key="manager_modifica"
        )
        
        manager_id = opzioni_manager[
            scelta_manager
        ]
        
        dettagli_manager = pd.read_sql(
            """
            SELECT *
            FROM utenti
            WHERE id = ?
            """,
            conn,
            params=(manager_id,)
        ).iloc[0]

        manager_principale = (
            dettagli_manager["username"]
            ==
            st.secrets.get(
                "MANAGER_USERNAME",
                ""
            )
        )

        st.text_input(
            "Email",
            value=dettagli_manager["username"],
            disabled=True
        )
        
        st.text_input(
            "Password corrente",
            value=dettagli_manager["password_visibile"]
                if pd.notna(
                    dettagli_manager["password_visibile"]
                )
                else "",
            disabled=True
        )

        if st.button(
            "📧 Reinvia credenziali manager"
        ):
        
            invia_credenziali_manager_email(
                dettagli_manager["username"],
                dettagli_manager["nome"],
                dettagli_manager["password_visibile"]
            )
        
            st.success(
                "Credenziali inviate."
            )

        nuova_password_manager = st.text_input(
            "Nuova password manager",
            type="password",
            key="password_manager"
        )

        if st.button(
            "🔐 Cambia password manager"
        ):
        
            if nuova_password_manager.strip() == "":
        
                st.error(
                    "Inserisci una password."
                )
        
            else:
        
                aggiorna_password_utente(
                    manager_id,
                    nuova_password_manager.strip()
                )
        
                st.success(
                    "Password aggiornata."
                )
        
                st.rerun()

        st.markdown("---")

        st.subheader("🗑️ Elimina manager")   

        if manager_principale:

            st.warning(
                "Il manager principale non può essere eliminato."
            )
        
        else:
        
            conferma_elimina_manager = st.checkbox(
                f"Confermo eliminazione di {dettagli_manager['nome']}",
                key=f"elimina_manager_{manager_id}"
            )
        
            if st.button(
                "🗑️ Elimina manager"
            ):
        
                if not conferma_elimina_manager:
        
                    st.error(
                        "Devi confermare."
                    )
        
                else:
        
                    elimina_manager(
                        manager_id
                    )
        
                    st.success(
                        "Manager eliminato."
                    )
        
                    st.rerun()


if is_manager():
    with tab_riepilogo:
            
        stagioni = get_stagioni()
    
        stagione_riepilogo = st.selectbox(
            "Stagione",
            stagioni,
            key="stagione_riepilogo"
        )
        
        st.header("📊 Riepilogo Presenze")
            
        bambini = get_bambini(
            attivi_solo=False
        )
            
        if bambini.empty:
            
            st.info(
                "Nessun bambino presente."
            )
            
        else:
            
            opzioni_bambini = {
                f"{row['cognome']} {row['nome']}": int(row["id"])
                for _, row in bambini.iterrows()
            }
            
            bambino_label = st.selectbox(
                "Seleziona bambino",
                list(opzioni_bambini.keys())
            )
            
            bambino_id = opzioni_bambini[
                bambino_label
            ]
            
            dati_bambino = bambini[
                bambini["id"] == bambino_id
            ].iloc[0]
            
            corso_id = int(
                dati_bambino["corso_id"]
            )
    
            corso_info = get_corso_by_id(
                corso_id
            )
                    
            nome_corso = (
                corso_info.iloc[0]["nome"]
            )
    
            giorni_corso = get_giorni_corso(
                corso_id
            )
            
            giorni_validi = (
                giorni_corso["giorno"]
                .tolist()
            )
            
            presenze = get_presenze_bambino(
                bambino_id
            )
            
            presenze_dict = {}
    
            for _, row in presenze.iterrows():
    
                presenze_dict[
                    str(row["data"])
                ] = int(
                    row["presenza"]
                )
    
            presenze_totali = sum(
                1
                for valore in presenze_dict.values()
                if valore == 1
            )
    
            assenze_totali = sum(
                1
                for valore in presenze_dict.values()
                if valore == 0
            )
    
            lezioni_programmate = (
                presenze_totali
                +
                assenze_totali
            )
    
            if lezioni_programmate > 0:
    
                percentuale_presenza = round(
                    presenze_totali
                    * 100
                    / lezioni_programmate,
                    2
                )
    
            else:
    
                percentuale_presenza = 0
    
            chiusure_df = get_chiusure()
    
            chiusure = set(
                chiusure_df["data"].tolist()
            )
    
            mappa_giorni = {
                "Lunedì": 0,
                "Martedì": 1,
                "Mercoledì": 2,
                "Giovedì": 3,
                "Venerdì": 4,
                "Sabato": 5,
                "Domenica": 6
            }
    
            stagione = stagione_riepilogo
    
            anno_inizio = int(
                stagione.split("/")[0]
            )
    
            anno_fine = anno_inizio + 1
    
            mesi = [
                ("Settembre", 9, anno_inizio),
                ("Ottobre", 10, anno_inizio),
                ("Novembre", 11, anno_inizio),
                ("Dicembre", 12, anno_inizio),
                ("Gennaio", 1, anno_fine),
                ("Febbraio", 2, anno_fine),
                ("Marzo", 3, anno_fine),
                ("Aprile", 4, anno_fine),
                ("Maggio", 5, anno_fine),
                ("Giugno", 6, anno_fine),
                ("Luglio", 7, anno_fine),
                ("Agosto", 8, anno_fine)
            ]
    
            festivita_fisse = {
                "01-01",
                "01-06",
                "04-25",
                "05-01",
                "06-02",
                "08-15",
                "11-01",
                "12-08",
                "12-25",
                "12-26"
            }
    
            tabella = []
    
            for nome_mese, mese, anno in mesi:
    
                riga = {
                    "Mese": nome_mese
                }
    
                giorni_mese = monthrange(
                    anno,
                    mese
                )[1]
    
                for giorno in range(1, 32):
    
                    if giorno > giorni_mese:
    
                        riga[str(giorno)] = ""
    
                        continue
    
                    data_corrente = datetime(
                        anno,
                        mese,
                        giorno
                    )
    
                    data_str = data_corrente.strftime(
                        "%Y-%m-%d"
                    )
    
                    corso_previsto = (
                        data_corrente.weekday()
                        in
                        [
                            mappa_giorni[g]
                            for g in giorni_validi
                        ]
                    )
    
                    if data_str in chiusure:
    
                        riga[str(giorno)] = "-"
    
                    elif (
                        data_corrente.strftime("%m-%d")
                        in festivita_fisse
                    ):
    
                        riga[str(giorno)] = "-"
    
                    elif not corso_previsto:
    
                        riga[str(giorno)] = "-"
    
                    elif data_str in presenze_dict:
    
                        if presenze_dict[data_str] == 1:
    
                            riga[str(giorno)] = "✅"
    
                        else:
    
                            riga[str(giorno)] = "❌"
    
                    else:
    
                        riga[str(giorno)] = ""
    
                tabella.append(
                    riga
                )
    
            df_calendario = pd.DataFrame(
                tabella
            )
    
            data_generazione = datetime.now().strftime(
                "%d/%m/%Y"
            )
    
            st.dataframe(
                df_calendario,
                use_container_width=True,
                hide_index=True
            )
    
            pdf_data = genera_pdf_presenze(
                bambino_label,
                nome_corso,
                stagione,
                presenze_totali,
                assenze_totali,
                percentuale_presenza,
                data_generazione,
                df_calendario
            )
    
            st.download_button(
                "📄 Scarica PDF",
                pdf_data,
                file_name=f"presenze_{bambino_label}.pdf",
                mime="application/pdf",
                key=f"pdf_presenze_bambino_{bambino_id}"
            )
    
            st.markdown("---")
            st.subheader("📅 Riepilogo mensile per corso")
                    
            corsi_riepilogo = get_corsi_con_giorni(
                attivi_solo=True,
                stagione=stagione_riepilogo
            )
                    
            if corsi_riepilogo.empty:
                    
                st.info(
                    "Nessun corso disponibile per questa stagione."
                )
                    
            else:
                    
                opzioni_corsi_mensile = {
                    f"{row['nome']} | {row['giorni_orari']}": int(row["id"])
                    for _, row in corsi_riepilogo.iterrows()
                }
                    
                corso_label_mensile = st.selectbox(
                    "Seleziona corso",
                    list(opzioni_corsi_mensile.keys()),
                    key="corso_riepilogo_mensile"
                )
                    
                corso_id_mensile = opzioni_corsi_mensile[
                    corso_label_mensile
                ]
                    
                dati_corso_mensile = get_corso_by_id(
                    corso_id_mensile
                )
                    
                nome_corso_mensile = dati_corso_mensile.iloc[0]["nome"]
                    
                mesi_riepilogo = [
                    ("Settembre", 9),
                    ("Ottobre", 10),
                    ("Novembre", 11),
                    ("Dicembre", 12),
                    ("Gennaio", 1),
                    ("Febbraio", 2),
                    ("Marzo", 3),
                    ("Aprile", 4),
                    ("Maggio", 5),
                    ("Giugno", 6),
                    ("Luglio", 7),
                    ("Agosto", 8)
                ]
                    
                opzioni_mesi = {
                    nome: numero
                    for nome, numero in mesi_riepilogo
                }
                    
                mese_label = st.selectbox(
                    "Seleziona mese",
                    list(opzioni_mesi.keys()),
                    key="mese_riepilogo_mensile"
                )
                    
                mese_numero = opzioni_mesi[
                    mese_label
                ]
                    
                anno_inizio = int(
                    stagione_riepilogo.split("/")[0]
                )
    
                if mese_numero >= 9:
    
                    anno_mese = anno_inizio
    
                else:
    
                    anno_mese = anno_inizio + 1
        
                giorni_corso_mensile = get_giorni_corso(
                    corso_id_mensile
                )
            
                giorni_validi_mensile = (
                    giorni_corso_mensile["giorno"]
                    .tolist()
                )
            
                mappa_giorni = {
                    "Lunedì": 0,
                    "Martedì": 1,
                    "Mercoledì": 2,
                    "Giovedì": 3,
                    "Venerdì": 4,
                    "Sabato": 5,
                    "Domenica": 6
                }
            
                bambini_corso = get_bambini_corso(
                    corso_id_mensile,
                    attivi_solo=True
                )
            
                if bambini_corso.empty:
            
                    st.info(
                        "Nessun bambino presente in questo corso."
                    )
            
                else:
            
                    presenze_corso_mese = pd.read_sql(
                        """
                        SELECT *
                        FROM presenze
                        WHERE corso_id = ?
                        AND data >= ?
                        AND data <= ?
                        """,
                        conn,
                        params=(
                            corso_id_mensile,
                            f"{anno_mese}-{mese_numero:02d}-01",
                            f"{anno_mese}-{mese_numero:02d}-{monthrange(anno_mese, mese_numero)[1]:02d}"
                        )
                    )
            
                    presenze_mensili_dict = {}
            
                    for _, row in presenze_corso_mese.iterrows():
            
                        chiave = (
                            int(row["bambino_id"]),
                            str(row["data"])
                        )
            
                        presenze_mensili_dict[
                            chiave
                        ] = int(
                            row["presenza"]
                        )
            
                    chiusure_df = get_chiusure()
            
                    chiusure = set(
                        chiusure_df["data"].tolist()
                    )
            
                    festivita_fisse = {
                        "01-01",
                        "01-06",
                        "04-25",
                        "05-01",
                        "06-02",
                        "08-15",
                        "11-01",
                        "12-08",
                        "12-25",
                        "12-26"
                    }
            
                    giorni_mese = monthrange(
                        anno_mese,
                        mese_numero
                    )[1]
                        
                    righe_mensili = []
        
                    for _, bambino in bambini_corso.iterrows():
            
                        bambino_id_mensile = int(
                            bambino["id"]
                        )
            
                        nome_bambino_mensile = (
                            f"{bambino['cognome']} {bambino['nome']}"
                        )
            
                        riga = {
                            "Bambino": nome_bambino_mensile
                        }
            
                        totale_presenze = 0
                        totale_assenze = 0
            
                        for giorno in range(1, 32):
            
                            if giorno > giorni_mese:
            
                                riga[str(giorno)] = ""
            
                                continue
            
                            data_corrente = datetime(
                                anno_mese,
                                mese_numero,
                                giorno
                            )
            
                            data_str = data_corrente.strftime(
                                "%Y-%m-%d"
                            )
            
                            corso_previsto = (
                                data_corrente.weekday()
                                in
                                [
                                    mappa_giorni[g]
                                    for g in giorni_validi_mensile
                                ]
                            )
            
                            if data_str in chiusure:
            
                                riga[str(giorno)] = "-"
            
                            elif (
                                data_corrente.strftime("%m-%d")
                                in festivita_fisse
                            ):
            
                                riga[str(giorno)] = "-"
        
                            elif not corso_previsto:
            
                                riga[str(giorno)] = "-"
        
                            elif (
                                 bambino_id_mensile,
                                data_str
                            ) in presenze_mensili_dict:
            
                                valore = presenze_mensili_dict[
                                    (
                                        bambino_id_mensile,
                                        data_str
                                    )
                                ]
            
                                if valore == 1:
            
                                    riga[str(giorno)] = "P"
                                    totale_presenze += 1
            
                                else:
            
                                    riga[str(giorno)] = "A"
                                    totale_assenze += 1
        
                            else:
            
                                riga[str(giorno)] = ""
            
                        riga["Presenze"] = totale_presenze
                        riga["Assenze"] = totale_assenze
            
                        lezioni_compilate = (
                            totale_presenze +
                            totale_assenze
                        )
            
                        if lezioni_compilate > 0:
            
                            riga["% Presenza"] = round(
                                totale_presenze
                                * 100
                                / lezioni_compilate,
                                2
                            )
            
                        else:
            
                            riga["% Presenza"] = 0
            
                        righe_mensili.append(
                            riga
                        )
                        
                    df_mensile = pd.DataFrame(
                        righe_mensili
                    )
                        
                    st.dataframe(
                        df_mensile,
                        use_container_width=True,
                        hide_index=True
                    )
                        
                    data_generazione_mensile = datetime.now().strftime(
                        "%d/%m/%Y"
                    )
                        
                    pdf_mensile = genera_pdf_riepilogo_mensile_corso(
                        nome_corso_mensile,
                        mese_label,
                        anno_mese,
                        stagione_riepilogo,
                        df_mensile,
                        data_generazione_mensile
                    )
                        
                    nome_file_corso = nome_corso_mensile.replace(
                        " ",
                        "_"
                    )
                        
                    st.download_button(
                        "📄 Scarica PDF riepilogo mensile corso",
                        pdf_mensile,
                        file_name=f"riepilogo_mensile_{nome_file_corso}_{mese_label}_{anno_mese}.pdf",
                        mime="application/pdf"
                    )

if is_manager():
    with tab_chiusure:

        st.header("🚫 Chiusure")
    
        data_chiusura = st.date_input(
            "Data"
        )
    
        descrizione = st.text_input(
            "Descrizione"
        )
    
        if st.button("➕ Inserisci chiusura"):
    
            aggiungi_chiusura(
                str(data_chiusura),
                descrizione
            )
    
            st.rerun()
    
        chiusure_df = get_chiusure()
    
        st.dataframe(
            chiusure_df,
            use_container_width=True
        )
    
        if not chiusure_df.empty:
        
            opzioni = {
                f"{row['data']} - {row['descrizione']}":
                int(row["id"])
                for _, row in chiusure_df.iterrows()
            }
        
            chiusura_label = st.selectbox(
                "Chiusura da eliminare",
                list(opzioni.keys())
            )
        
            if st.button(
                "🗑️ Elimina chiusura"
            ):
        
                elimina_chiusura(
                    opzioni[chiusura_label]
                )
        
                st.success(
                    "Chiusura eliminata."
                )
        
                st.rerun()

if is_manager():
    with tab_genitori:

        st.header(
            "👨‍👩‍👧 Gestione Genitori"
        )
    
        genitori = get_genitori()
    
        if genitori.empty:
    
            st.info(
                "Nessun genitore presente."
            )
    
        else:
    
            st.dataframe(
                genitori[
                    [
                        "nome",
                        "username",
                        "telefono_genitore",
                        "nome_bambino",
                        "cognome_bambino",
                        "attivo"
                    ]
                ].rename(
                    columns={
                        "nome": "Genitore",
                        "username": "Email",
                        "telefono_genitore": "Telefono",
                        "nome_bambino": "Nome bambino",
                        "cognome_bambino": "Cognome bambino",
                        "attivo": "Attivo"
                    }
                ),
                use_container_width=True,
                hide_index=True
            )
            
            opzioni = {
                f"{r['nome']} ({r['username']})":
                int(r["id"])
                for _, r in genitori.iterrows()
            }
    
            genitore_label = st.selectbox(
                "Genitore",
                list(opzioni.keys())
            )
    
            genitore_id = opzioni[
                genitore_label
            ]
    
            dettagli = pd.read_sql(
                """
                SELECT *
                FROM utenti
                WHERE id = ?
                """,
                conn,
                params=(genitore_id,)
            ).iloc[0]
    
            bambino_assoc = (
                get_bambino_associato_genitore(
                    genitore_id
                )
            )

            bambino_id_df = pd.read_sql(
                """
                SELECT bambino_id
                FROM genitori_bambini
                WHERE utente_id = ?
                """,
                conn,
                params=(genitore_id,)
            )
            
            if bambino_id_df.empty:
            
                st.error(
                    "Nessun bambino associato."
                )
            
                st.stop()
            
            bambino_id = int(
                bambino_id_df.iloc[0]["bambino_id"]
            )

            email = pd.read_sql(
                """
                SELECT u.username
                FROM utenti u
                WHERE u.id = ?
                """,
                conn,
                params=(genitore_id,)
            )
            
            if not email.empty:
            
                st.text_input(
                    "Email genitore",
                    value=email.iloc[0]["username"]
                    if pd.notna(email.iloc[0]["username"])
                    else "",
                    disabled=True,
                    key=f"email_genitore_visualizzazione_{genitore_id}"
                )

            telefono = pd.read_sql(
                    """
                    SELECT b.telefono_genitore
                    FROM bambini b
                    JOIN genitori_bambini gb
                        ON gb.bambino_id = b.id
                    WHERE gb.utente_id = ?
                    """,
                    conn,
                    params=(genitore_id,)
                )
                
            if not telefono.empty:
                
                st.text_input(
                    "Telefono genitore",
                    value=telefono.iloc[0]["telefono_genitore"]
                    if pd.notna(telefono.iloc[0]["telefono_genitore"])
                    else "",
                    disabled=True,
                    key="telefono_genitore_visualizzazione"
                )
    
            if not bambino_assoc.empty:
    
                st.info(
                    f"Bambino associato: "
                    f"{bambino_assoc.iloc[0]['cognome']} "
                    f"{bambino_assoc.iloc[0]['nome']}"
                )
            
            scheda = pd.read_sql(
                """
                SELECT *
                FROM schede_genitori
                WHERE bambino_id = ?
                """,
                conn,
                params=(
                    int(
                        bambino_id
                    ),
                )
            )

            if scheda.empty:
                
                st.warning(
                    "Il genitore non ha ancora inviato la pratica."
                )

            else:

                scheda = scheda.iloc[0]

                st.subheader(
                    "👦 Dati atleta"
                )
                
                st.text_input(
                    "Luogo nascita",
                    value=scheda["luogo_nascita_bambino"]
                    if pd.notna(
                        scheda["luogo_nascita_bambino"]
                    )
                    else ""
                )
                
                st.text_input(
                    "Data nascita",
                    value=scheda["data_nascita_bambino"]
                    if pd.notna(
                        scheda["data_nascita_bambino"]
                    )
                    else ""
                )
                
                st.text_input(
                    "Codice fiscale",
                    value=scheda["cf_bambino"]
                    if pd.notna(
                        scheda["cf_bambino"]
                    )
                    else ""
                )
                
                st.text_input(
                    "Comune",
                    value=scheda["comune_bambino"]
                    if pd.notna(
                        scheda["comune_bambino"]
                    )
                    else ""
                )
                
                st.text_input(
                    "Indirizzo",
                    value=scheda["indirizzo_bambino"]
                    if pd.notna(
                        scheda["indirizzo_bambino"]
                    )
                    else ""
                )

                st.subheader(
                    "👨‍👩‍👧 Dati genitore"
                )
                
                nome_genitore = st.text_input(
                    "Nome genitore",
                    value=scheda["nome_genitore"]
                    if pd.notna(
                        scheda["nome_genitore"]
                    )
                    else ""
                )
                
                cognome_genitore = st.text_input(
                    "Cognome genitore",
                    value=scheda["cognome_genitore"]
                    if pd.notna(
                        scheda["cognome_genitore"]
                    )
                    else ""
                )
                
                cf_genitore = st.text_input(
                    "Codice fiscale genitore",
                    value=scheda["cf_genitore"]
                    if pd.notna(
                        scheda["cf_genitore"]
                    )
                    else ""
                )

                st.subheader(
                    "📋 Stato pratica"
                )
                
                if scheda["pratica_inviata"] == 1:
                
                    st.success(
                        f"""
                        ✅ Pratica inviata
                
                        Data invio:
                        {scheda["data_invio"]}
                        """
                    )
                
                else:
                
                    st.warning(
                        "Pratica non inviata."
                    )

                if st.button(
                    "💾 Salva modifiche manager"
                ):

                    c.execute(
                        """
                        UPDATE schede_genitori
                        SET
                    
                            nome_genitore = ?,
                            cognome_genitore = ?,
                            cf_genitore = ?
                    
                        WHERE bambino_id = ?
                        """,
                        (
                            nome_genitore,
                            cognome_genitore,
                            cf_genitore,
                            int(
                                bambino_id
                            )
                        )
                    )
                    
                    conn.commit()
                    
                    st.success(
                        "Modifiche salvate."
                    )

                if st.button(
                    "🔓 Riapri pratica al genitore"
                ):
                    c.execute(
                        """
                        UPDATE schede_genitori
                        SET
                    
                            bloccato = 0,
                            pratica_inviata = 0
                    
                        WHERE bambino_id = ?
                        """,
                        (
                            int(
                                bambino_id
                            ),
                        )
                    )
                    
                    conn.commit()
                    
                    st.success(
                        "Pratica riaperta."
                    )
                    
                    st.rerun()


############## AREA PERSONALE #################

if is_genitore():

    with tab_area_personale:
        bambino = get_bambino_genitore(
            st.session_state["utente_id"]
        )
        
        if bambino.empty:
        
            st.warning(
                "Nessun bambino associato."
            )
        
            st.stop()

        bambino_id = int(
            bambino.iloc[0]["id"]
        )

        st.subheader("👶 Dati Bambino/a")

        scheda = pd.read_sql(
               """
               SELECT *
               FROM schede_genitori
               WHERE bambino_id = ?
               """,
               conn,
               params=(bambino_id,)
            )
    
        bloccato = False
    
        if (
            not scheda.empty
            and
            "bloccato" in scheda.columns
            and
            int(scheda.iloc[0]["bloccato"]) == 1
        ):
            bloccato = True
                
        valori = {}
        
        if not scheda.empty:
            
            valori = scheda.iloc[0].to_dict()

        col1, col2 = st.columns(2)
        
        with col1:
        
            st.text_input(
                "Nome",
                value=bambino.iloc[0]["nome"],
                disabled=True
            )
        
            luogo_nascita_bambino = st.text_input(
                "Luogo di nascita del bambino",
                value=valori.get(
                    "luogo_nascita_bambino",
                    ""
                ),
                disabled=bloccato,
                key=f"luogo_nascita_bambino_{bambino_id}"
            )

            cf_bambino = st.text_input(
                "Codice fiscale del bambino",
                value=valori.get(
                    "cf_bambino",
                    ""
                ),
                disabled=bloccato,
                key=f"cf_bambino_{bambino_id}"
            )

            indirizzo_bambino = st.text_input(
                "Indirizzo di residenza del bambino",
                value=valori.get(
                    "indirizzo_bambino",
                    ""
                ),
                disabled=bloccato,
                key=f"indirizzo_bambino_{bambino_id}"
            )

            st.text_input(
                "Email genitore",
                value=bambino.iloc[0]["email_genitore"]
                if pd.notna(bambino.iloc[0]["email_genitore"])
                else "",
                disabled=True
            )

            corso = pd.DataFrame()

            giorni = pd.DataFrame()
    
            if pd.notna(bambino.iloc[0]["corso_id"]):
            
                corso = pd.read_sql(
                    """
                    SELECT
                        nome,
                        stagione
                    FROM corsi
                    WHERE id = ?
                    """,
                    conn,
                    params=(
                        int(bambino.iloc[0]["corso_id"]),
                    )
                )
            
                giorni = pd.read_sql(
                    """
                    SELECT
                        giorno,
                        orario
                    FROM corso_giorni
                    WHERE corso_id = ?
                    ORDER BY ordine
                    """,
                    conn,
                    params=(
                        int(bambino.iloc[0]["corso_id"]),
                    )
                )
                        
                with col2:

                    st.text_input(
                        "Cognome",
                        value=bambino.iloc[0]["cognome"],
                        disabled=True
                    ) 

                    valore_data = date(2018, 1, 1)

                    if not scheda.empty and valori.get("data_nascita_bambino"):
                    
                        valore_data = datetime.strptime(
                            valori["data_nascita_bambino"],
                            "%Y-%m-%d"
                        ).date()
                    
                    data_nascita_bambino = st.date_input(
                        "Data di nascita del bambino",
                        value=valore_data,
                        min_value=date(2000, 1, 1),
                        max_value=date.today(),
                        disabled=bloccato,
                        format="DD/MM/YYYY",
                        key=f"data_nascita_bambino_{bambino_id}"
                    )

                    eta_bambino = calcola_eta(
                        data_nascita_bambino
                    )
                    
                    richiede_certificato_medico = (
                        eta_bambino >= 6
                    )

                    comune_bambino = st.text_input(
                        "Comune di residenza",
                        value=valori.get(
                            "comune_bambino",
                            ""
                        ),
                        disabled=bloccato,
                        key=f"comune_bambino_{bambino_id}"
                    )

                    cap_bambino = st.text_input(
                        "CAP",
                        value=valori.get(
                            "cap_bambino",
                            ""
                        ),
                        disabled=bloccato,
                        key=f"cap_bambino_{bambino_id}"
                    )
                
                    st.text_input(
                        "Telefono genitore",
                        value=bambino.iloc[0]["telefono_genitore"]
                        if pd.notna(bambino.iloc[0]["telefono_genitore"])
                        else "",
                        disabled=True
                    )

                    st.text_input(
                        "Stagione sportiva",
                        value=corso.iloc[0]["stagione"],
                        disabled=True,
                        key=f"stagione_corso_{bambino_id}"
                    )
        
                if not corso.empty:
        
                    st.text_input(
                        "Corso",
                        value=corso.iloc[0]["nome"],
                        disabled=True,
                        key=f"nome_corso_{bambino_id}"
                    )

                    st.dataframe(
                        giorni,
                        hide_index=True,
                        width="stretch"
                    )
                
        st.markdown("---")
        
        st.subheader("👨‍👩‍👧 Dati genitore")
        
        col1, col2 = st.columns(2)
        
        with col1:
            nome_genitore = st.text_input(
                "Nome genitore",
                value=valori.get("nome_genitore", ""),
                disabled=bloccato
            )
        
        with col2:
            cognome_genitore = st.text_input(
                "Cognome genitore",
                value=valori.get("cognome_genitore", ""),
                disabled=bloccato
            )
        
        col1, col2 = st.columns(2)
        
        with col1:
            luogo_nascita_genitore = st.text_input(
                "Luogo di nascita",
                value=valori.get("luogo_nascita_genitore", ""),
                disabled=bloccato
            )
        
        with col2:

            valore_data = date(1990, 1, 1)

            if (
                not scheda.empty
                and valori.get("data_nascita_genitore")
            ):
            
                try:
            
                    valore_data = datetime.strptime(
                        valori["data_nascita_genitore"],
                        "%Y-%m-%d"
                    ).date()
            
                except:
            
                    pass
            
            data_nascita_genitore = st.date_input(
                "Data di nascita del genitore",
                value=valore_data,
                min_value=date(1940, 1, 1),
                max_value=date.today(),
                disabled=bloccato,
                format="DD/MM/YYYY",
                key=f"data_nascita_genitore_{bambino_id}"
            )
        
        col1, col2 = st.columns(2)
        
        with col1:
            cf_genitore = st.text_input(
                "Codice fiscale",
                value=valori.get("cf_genitore", ""),
                disabled=bloccato
            )
        
        with col2:
            cap_genitore = st.text_input(
                "CAP",
                value=valori.get("cap_genitore", ""),
                disabled=bloccato
            )
        
        col1, col2 = st.columns(2)
        
        with col1:
            indirizzo_genitore = st.text_input(
                "Indirizzo",
                value=valori.get("indirizzo_genitore", ""),
                disabled=bloccato
            )
        
        with col2:
            comune_genitore = st.text_input(
                "Comune",
                value=valori.get("comune_genitore", ""),
                disabled=bloccato
            )
        
        
        if st.button(
            "💾 Salva anagrafica",
            disabled=bloccato
        ):
        
            campi_mancanti = []

            controlli = {
            
                "Luogo nascita bambino": luogo_nascita_bambino,
                "CF bambino": cf_bambino,
                "Indirizzo bambino": indirizzo_bambino,
                "Comune bambino": comune_bambino,
                "CAP bambino": cap_bambino,
            
                "Nome genitore": nome_genitore,
                "Cognome genitore": cognome_genitore,
                "Luogo nascita genitore": luogo_nascita_genitore,
                "CF genitore": cf_genitore,
                "Indirizzo genitore": indirizzo_genitore,
                "Comune genitore": comune_genitore,
                "CAP genitore": cap_genitore
            
            }
            
            for nome, valore in controlli.items():
            
                if str(valore).strip() == "":
                    campi_mancanti.append(nome)

                if campi_mancanti:
    
                    st.error(
                        "Campi mancanti:\n\n- "
                        + "\n- ".join(campi_mancanti)
                    )
            
                # INSERT / UPDATE
        
            else:
        
                if scheda.empty:
        
                    c.execute(
                        """
                        INSERT INTO schede_genitori (

                            bambino_id,
                            data_nascita_bambino,
                        
                            luogo_nascita_bambino,
                            cf_bambino,
                            indirizzo_bambino,
                            comune_bambino,
                            cap_bambino,
                        
                            nome_genitore,
                            cognome_genitore,
        
                            luogo_nascita_genitore,
                            data_nascita_genitore,
        
                            cf_genitore,
        
                            indirizzo_genitore,
                            comune_genitore,
                            cap_genitore,
        
                            ultima_modifica
        
                        )
        
                        VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                        """,
                        (
                            bambino_id,
                            str(data_nascita_bambino),
                        
                            luogo_nascita_bambino,
                            cf_bambino,
                            indirizzo_bambino,
                            comune_bambino,
                            cap_bambino,
                        
                            nome_genitore,
                            cognome_genitore,
        
                            luogo_nascita_genitore,
                            data_nascita_genitore,
        
                            cf_genitore,
        
                            indirizzo_genitore,
                            comune_genitore,
                            cap_genitore,
        
                            str(datetime.now())
                        )
                    )
        
                else:
        
                    c.execute(
                        """
                        UPDATE schede_genitori
                        SET

                            data_nascita_bambino = ?,
                            
                            luogo_nascita_bambino = ?,
                            cf_bambino = ?,
                            indirizzo_bambino = ?,
                            comune_bambino = ?,
                            cap_bambino = ?,
                            
                            nome_genitore = ?,
                            cognome_genitore = ?,
        
                            luogo_nascita_genitore = ?,
                            data_nascita_genitore = ?,
        
                            cf_genitore = ?,
        
                            indirizzo_genitore = ?,
                            comune_genitore = ?,
                            cap_genitore = ?,
        
                            ultima_modifica = ?
        
                        WHERE bambino_id = ?
                        """,
                        (
                            str(data_nascita_bambino),
                        
                            luogo_nascita_bambino,
                            cf_bambino,
                            indirizzo_bambino,
                            comune_bambino,
                            cap_bambino,
                        
                            nome_genitore,
                            cognome_genitore,
        
                            luogo_nascita_genitore,
                            data_nascita_genitore,
        
                            cf_genitore,
        
                            indirizzo_genitore,
                            comune_genitore,
                            cap_genitore,
        
                            str(datetime.now()),
        
                            bambino_id
                        )
                    )
        
                conn.commit()
        
                st.success(
                    "Anagrafica salvata."
                )
        
                st.rerun()

        st.markdown("---")

        st.subheader("📄 Modulistica")

        scheda = pd.read_sql(
            """
            SELECT *
            FROM schede_genitori
            WHERE bambino_id = ?
            """,
            conn,
            params=(bambino_id,)
        )

        anagrafica_completa = False

        if not scheda.empty:
        
            r = scheda.iloc[0]
        
            anagrafica_completa = all(
                [
            
                    # DATI BAMBINO
            
                    str(r.get("data_nascita_bambino", "")).strip(),
                    str(r.get("luogo_nascita_bambino", "")).strip(),
                    str(r.get("cf_bambino", "")).strip(),
                    str(r.get("indirizzo_bambino", "")).strip(),
                    str(r.get("comune_bambino", "")).strip(),
                    str(r.get("cap_bambino", "")).strip(),
            
                    # DATI GENITORE
            
                    str(r.get("nome_genitore", "")).strip(),
                    str(r.get("cognome_genitore", "")).strip(),
                    str(r.get("luogo_nascita_genitore", "")).strip(),
                    str(r.get("data_nascita_genitore", "")).strip(),
                    str(r.get("cf_genitore", "")).strip(),
                    str(r.get("indirizzo_genitore", "")).strip(),
                    str(r.get("comune_genitore", "")).strip(),
                    str(r.get("cap_genitore", "")).strip(),
            
                ]
            )

        if not anagrafica_completa:

            st.warning(
                """
                Prima di poter generare
                la modulistica devi
                completare e salvare
                l'anagrafica.
                """
            )
        
        else:

            st.success(
                """
                Anagrafica completa.
                La modulistica può
                essere generata.
                """
            )
            
            if st.button(
                "📄 Genera modulo",
                key=f"genera_modulo_docx_{bambino_id}"
            ):
            
                try:
            
                    template_id = st.secrets[
                        "DRIVE_TEMPLATE_MODULO_ID"
                    ]
            
                    template_bytes = scarica_file_drive_bytes(
                        template_id
                    )
            
                    mappa = crea_mappa_modulo(
                        bambino,
                        scheda
                    )
            
                    docx_compilato = compila_docx_template(
                        template_bytes,
                        mappa
                    )
            
                    st.session_state[
                        f"docx_modulo_{bambino_id}"
                    ] = docx_compilato.getvalue()

                    pdf_compilato = genera_pdf_compilato(
                        bambino,
                        scheda
                    )
                    
                    st.session_state[
                        f"pdf_modulo_{bambino_id}"
                    ] = pdf_compilato.getvalue()

                    st.success(
                        "Modulo generato correttamente."
                    )
            
                except Exception as e:
            
                    st.error(
                        "Errore durante la generazione del modulo."
                    )
            
                    st.code(
                        str(e)
                    )
            
            # DOWNLOAD DOCX

            #if f"docx_modulo_{bambino_id}" in st.session_state:
            
             #   st.download_button(
              #      "📥 Scarica modulo DOCX",
               #     data=st.session_state[
                #        f"docx_modulo_{bambino_id}"
                 #   ],
                  #  file_name=(
                   #     f"Modulo_{bambino.iloc[0]['cognome']}_"
                    #    f"{bambino.iloc[0]['nome']}.docx"
                    #),
                    #mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                #)
            
            
            # DOWNLOAD PDF
            
            if f"pdf_modulo_{bambino_id}" in st.session_state:
            
                st.download_button(
                    "📥 Scarica modulo PDF",
                    data=st.session_state[
                        f"pdf_modulo_{bambino_id}"
                    ],
                    file_name=(
                        f"Modulo_{bambino.iloc[0]['cognome']}_"
                        f"{bambino.iloc[0]['nome']}.pdf"
                    ),
                    mime="application/pdf"
                )
                
        # ============================================================
        # INTERFACCIA
        # ============================================================

        st.markdown("---")

        st.subheader("📤 Documenti")

        modulo_firmato = st.file_uploader(
            "Modulo iscrizione firmato",
            type=["pdf"],
            disabled=bloccato,
            key=f"modulo_firmato_{bambino_id}"
        )

        if richiede_certificato_medico:

            st.info(
                f"Età bambino: {eta_bambino} anni. "
                "Il certificato medico è obbligatorio."
            )
        
            certificato_medico = st.file_uploader(
                "Certificato medico del bambino",
                type=["pdf"],
                disabled=bloccato,
                key=f"certificato_medico_{bambino_id}"
            )
        
        else:
        
            st.success(
                f"Età bambino: {eta_bambino} anni. "
                "Il certificato medico non è richiesto."
            )
        
            certificato_medico = True
    
        documento_identita = st.file_uploader(
            "Documento identità bambino",
            type=["pdf","jpg","jpeg","png"],
            disabled=bloccato,
            key=f"documento_identita_{bambino_id}"
        )

        tessera_sanitaria = st.file_uploader(
            "Tessera sanitaria bambino",
            type=["pdf","jpg","jpeg","png"],
            disabled=bloccato,
            key=f"tessera_sanitaria_{bambino_id}"
        )

        documento_identita_genitore = st.file_uploader(
            "Documento identità genitore",
            type=["pdf","jpg","jpeg","png"],
            disabled=bloccato,
            key=f"documento_identita_genitore_{bambino_id}"
        )

        st.info(
            """
            I documenti verranno
            successivamente salvati.
            """
        )

        st.markdown("---")

        st.subheader("✅ Invio pratica")

        pratica_inviata = (
            scheda.iloc[0].get(
                "pratica_inviata",
                0
            )
        )

        completa = pratica_completa(
            scheda,
            modulo_firmato,
            certificato_medico,
            documento_identita,
            tessera_sanitaria,
            documento_identita_genitore
        )

        if pratica_inviata == 1:

            st.success(
                f"""
                ✅ Pratica già inviata
        
                Data invio:
                {scheda.iloc[0]['data_invio']}
                """
            )
        
        else:
        
            completa = pratica_completa(
                scheda,
                modulo_firmato,
                certificato_medico,
                documento_identita,
                tessera_sanitaria,
                documento_identita_genitore
            )
        
            if not completa:
        
                st.error(
                    """
                    Documentazione incompleta.
        
                    Compila tutti i campi
                    e carica tutti i file.
                    """
                )
        
            else:
        
                st.success(
                    """
                    La documentazione
                    è completa.
        
                    Puoi inviare la pratica.
                    """
                )

                if (
                    completa
                    and
                    st.button(
                        "✅ Invia documentazione"
                    )
                ):
        
                    # ==========================
                    # CARTELLA STAGIONE
                    # ==========================
                    
                    nome_stagione = stagione_corrente()
        
                    invia_documentazione_email(
                        bambino,
                        modulo_firmato,
                        certificato_medico,
                        documento_identita,
                        tessera_sanitaria,
                        documento_identita_genitore
                    )
        
                    c.execute(
                        """
                        UPDATE schede_genitori
                        SET
                    
                            compilato = 1,
                    
                            bloccato = 1,
                    
                            pratica_inviata = 1,
                    
                            data_invio = ?
                    
                        WHERE bambino_id = ?
                        """,
                        (
                            str(datetime.now()),
                            bambino_id
                        )
                    )
                    
                    conn.commit()
        
                    pratica_inviata = scheda.iloc[0]["pratica_inviata"]
        
                    data_invio = scheda.iloc[0]["data_invio"]
        
                    if pratica_inviata:
                    
                        st.success(
                            f"""
                            ✅ Pratica inviata
                    
                            Data invio:
                            {data_invio}
                    
                            Il modulo è stato inviato correttamente.
                            La scheda è bloccata e non può più essere modificata.
                            """
                        )
    
            st.rerun()




